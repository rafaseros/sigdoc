import asyncio
import io
import os
import re
import tempfile

from docx import Document
from docxtpl import DocxTemplate

from app.domain.ports.template_engine import TemplateEngine


def _normalize_paragraph(paragraph: str) -> str:
    """
    Normalize a paragraph string for use as variable context.

    Returns the paragraph stripped of leading/trailing whitespace.
    Returning the full paragraph (rather than a truncated snippet) ensures
    that all variables found in the same paragraph share the SAME context
    string, so the frontend can deduplicate by exact string match.
    """
    return paragraph.strip()


def _camel_to_snake(name: str) -> str:
    """Convert camelCase/PascalCase to snake_case."""
    s1 = re.sub(r"(.)([A-Z][a-z]+)", r"\1_\2", name)
    return re.sub(r"([a-z0-9])([A-Z])", r"\1_\2", s1).lower()


class DocxTemplateEngine(TemplateEngine):
    """
    Template engine adapter using python-docx-template (docxtpl).

    CRITICAL DESIGN NOTES:
    - DocxTemplate CANNOT be reused after render() — internal XML is modified in-place.
      A FRESH instance must be created for every render call.
    - All operations run via asyncio.to_thread() to avoid blocking the event loop.
    - Always use autoescape=True to prevent XML injection from user-provided values.
    """

    async def extract_variables(self, file_bytes: bytes) -> list[dict]:
        """
        Extract all Jinja2 variable names AND their surrounding paragraph context
        from a .docx template.

        Returns: [{"name": "variable", "contexts": ["...paragraph text..."]}, ...]
        """

        def _extract(data: bytes) -> list[dict]:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp.write(data)
                tmp_path = tmp.name
            try:
                # Use DocxTemplate to get variable names (reliable extraction)
                tpl = DocxTemplate(tmp_path)
                variables = tpl.get_undeclared_template_variables()

                # Use python-docx to extract paragraph context
                doc = Document(tmp_path)
                all_paragraphs: list[str] = []

                # Collect paragraphs from body, headers, and footers
                for para in doc.paragraphs:
                    if para.text.strip():
                        all_paragraphs.append(para.text)
                for section in doc.sections:
                    if section.header:
                        for para in section.header.paragraphs:
                            if para.text.strip():
                                all_paragraphs.append(para.text)
                    if section.footer:
                        for para in section.footer.paragraphs:
                            if para.text.strip():
                                all_paragraphs.append(para.text)

                # Build variable -> contexts mapping, tracking order of first appearance
                var_contexts: dict[str, list[str]] = {}
                var_order: list[str] = []  # preserves first-appearance order
                var_pattern = re.compile(r"\{\{\s*(\w+)\s*\}\}")

                for para_text in all_paragraphs:
                    found_vars = var_pattern.findall(para_text)
                    for var_name in found_vars:
                        if var_name not in variables:
                            continue
                        if var_name not in var_contexts:
                            var_contexts[var_name] = []
                            var_order.append(var_name)
                        context = _normalize_paragraph(para_text)
                        if context not in var_contexts[var_name]:
                            var_contexts[var_name].append(context)

                # Add any variables found by docxtpl but not in paragraphs (e.g. in tables)
                for v in variables:
                    if v not in var_contexts:
                        var_contexts[v] = []
                        var_order.append(v)

                return [
                    {"name": name, "contexts": var_contexts[name]}
                    for name in var_order
                ]
            finally:
                os.unlink(tmp_path)

        return await asyncio.to_thread(_extract, file_bytes)

    async def render(self, file_bytes: bytes, variables: dict[str, str]) -> bytes:
        """
        Render a template with the given variables.

        IMPORTANT: Creates a FRESH DocxTemplate instance every time.
        Cannot reuse — render() modifies internal XML in-place.
        """

        def _render(data: bytes, context: dict[str, str]) -> bytes:
            # Write template to temp file
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp.write(data)
                tmp_path = tmp.name
            try:
                # FRESH instance — NEVER reuse after render
                tpl = DocxTemplate(tmp_path)
                tpl.render(context, autoescape=True)

                # Save rendered document to BytesIO
                output = io.BytesIO()
                tpl.save(output)
                output.seek(0)
                return output.read()
            finally:
                os.unlink(tmp_path)

        return await asyncio.to_thread(_render, file_bytes, variables)

    async def validate(self, file_bytes: bytes) -> dict:
        """
        Validate a template file and return validation results.

        Returns a dict with:
        - valid: bool
        - variables: list[str] — extracted variable names
        - errors: list of error dicts
        - has_fixable_errors: bool
        - has_unfixable_errors: bool
        """

        def _validate_sync(data: bytes) -> dict:
            from collections import Counter

            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp.write(data)
                tmp_path = tmp.name

            try:
                doc = Document(tmp_path)
                errors: list[dict] = []    # Block upload
                warnings: list[dict] = []  # Informational only
                variable_counts: Counter = Counter()
                variable_order: list[str] = []  # first-appearance order
                error_variables: set[str] = set()
                contexts_by_var: dict[str, list[str]] = {}  # var -> paragraph contexts

                # Collect all paragraphs: body + headers + footers
                all_paragraphs = list(doc.paragraphs)
                for section in doc.sections:
                    if section.header:
                        all_paragraphs.extend(section.header.paragraphs)
                    if section.footer:
                        all_paragraphs.extend(section.footer.paragraphs)

                for para in all_paragraphs:
                    full_text = para.text

                    # 1. No-space variables: {{var}} — FIXABLE (auto-fix adds spaces)
                    no_space_matches = re.findall(r"\{\{(\w+)\}\}", full_text)
                    for match in no_space_matches:
                        if match not in variable_counts:
                            variable_order.append(match)
                        variable_counts[match] += 1
                        # Check if the NAME itself is bad (uppercase, special chars)
                        if match != match.lower() or re.search(r"[A-Z]", match):
                            error_variables.add(match)
                            suggested = _camel_to_snake(match.lower())
                            errors.append({
                                "type": "uppercase",
                                "message": (
                                    f"Nombre de variable inválido: {{{{{match}}}}}. "
                                    f"Use {{{{ {suggested} }}}}"
                                ),
                                "variable": match,
                                "fixable": True,
                                "suggestion": suggested,
                            })
                        elif re.search(r"[^\x00-\x7F]", match):
                            error_variables.add(match)
                            errors.append({
                                "type": "special_chars",
                                "message": (
                                    f"Variable con caracteres especiales: {{{{{match}}}}}. "
                                    "Use solo letras ASCII, números y guiones bajos."
                                ),
                                "variable": match,
                                "fixable": False,
                                "suggestion": None,
                            })
                        else:
                            # Name is OK but missing spaces — just a warning
                            warnings.append({
                                "type": "no_spaces",
                                "message": (
                                    f"Variable sin espacios: {{{{{match}}}}}. "
                                    f"Se recomienda usar {{{{ {match} }}}}"
                                ),
                                "variable": match,
                                "fixable": True,
                                "suggestion": match,
                            })

                    # 2. Properly formatted variables {{ var }}
                    proper_matches = re.findall(r"\{\{\s+(\w+)\s+\}\}", full_text)
                    for match in proper_matches:
                        if match not in variable_counts:
                            variable_order.append(match)
                        variable_counts[match] += 1

                        # Check uppercase / camelCase — ERROR (blocks upload)
                        if match != match.lower() or re.search(r"[A-Z]", match):
                            error_variables.add(match)
                            suggested = _camel_to_snake(match.lower())
                            errors.append({
                                "type": "uppercase",
                                "message": (
                                    f"Nombre de variable inválido: {{{{ {match} }}}}. "
                                    f"Use {{{{ {suggested} }}}}"
                                ),
                                "variable": match,
                                "fixable": True,
                                "suggestion": suggested,
                            })

                        # Check special characters — ERROR (blocks upload)
                        elif re.search(r"[^\x00-\x7F]", match):
                            error_variables.add(match)
                            errors.append({
                                "type": "special_chars",
                                "message": (
                                    f"Variable con caracteres especiales: {{{{ {match} }}}}. "
                                    "Use solo letras ASCII, números y guiones bajos."
                                ),
                                "variable": match,
                                "fixable": False,
                                "suggestion": None,
                            })

                    # Collect paragraph context for every variable found in this paragraph
                    all_para_vars = re.findall(r"\{\{\s*(\w+)\s*\}\}", full_text)
                    for var_name in all_para_vars:
                        if var_name not in contexts_by_var:
                            contexts_by_var[var_name] = []
                        ctx = _normalize_paragraph(full_text)
                        if ctx not in contexts_by_var[var_name]:
                            contexts_by_var[var_name].append(ctx)

                    # 3. Split formatting — WARNING only (style issue, not blocking)
                    runs_text = [run.text for run in para.runs]
                    combined = "".join(runs_text)
                    combined_vars = re.findall(r"\{\{.*?\}\}", combined)
                    for cv in combined_vars:
                        found_in_single_run = any(cv in rt for rt in runs_text)
                        if not found_in_single_run:
                            var_name_match = re.search(r"\{\{\s*(\w+)\s*\}\}", cv)
                            name = var_name_match.group(1) if var_name_match else cv
                            warnings.append({
                                "type": "split_format",
                                "message": (
                                    f"La variable {{{{ {name} }}}} tiene formato dividido. "
                                    "Podría no funcionar correctamente. Se recomienda "
                                    "seleccionar todo el marcador en Word y aplicar el "
                                    "formato de manera uniforme."
                                ),
                                "variable": name,
                                "fixable": False,
                                "suggestion": None,
                            })

                    # 4. Unclosed braces — ERROR
                    if "{{" in full_text and "}}" not in full_text:
                        errors.append({
                            "type": "invalid_syntax",
                            "message": (
                                f'Llaves sin cerrar detectadas en: "{full_text[:80]}..."'
                            ),
                            "variable": None,
                            "fixable": False,
                            "suggestion": None,
                        })

                # 5. No variables at all — ERROR
                if not variable_counts and not errors:
                    errors.append({
                        "type": "no_variables",
                        "message": (
                            "No se encontraron variables en el documento. "
                            "Asegúrese de usar la sintaxis {{ nombre_variable }} con espacios."
                        ),
                        "variable": None,
                        "fixable": False,
                        "suggestion": None,
                    })

                # Deduplicate by (type, variable)
                def _dedup(items: list[dict]) -> list[dict]:
                    seen: set[tuple] = set()
                    result: list[dict] = []
                    for e in items:
                        key = (e["type"], e["variable"])
                        if key not in seen:
                            seen.add(key)
                            result.append(e)
                    return result

                unique_errors = _dedup(errors)
                unique_warnings = _dedup(warnings)

                has_fixable = any(e["fixable"] for e in unique_errors)
                has_unfixable = any(not e["fixable"] for e in unique_errors)

                # Build variable summary in document order
                variable_summary = []
                for var_name in variable_order:
                    variable_summary.append({
                        "name": var_name,
                        "count": variable_counts[var_name],
                        "has_errors": var_name in error_variables,
                        "contexts": contexts_by_var.get(var_name, []),
                    })

                return {
                    "valid": len(unique_errors) == 0,
                    "variables": [
                        v for v in variable_order if v not in error_variables
                    ],
                    "variable_summary": variable_summary,
                    "errors": unique_errors,
                    "warnings": unique_warnings,
                    "has_fixable_errors": has_fixable,
                    "has_unfixable_errors": has_unfixable,
                }
            finally:
                os.unlink(tmp_path)

        return await asyncio.to_thread(_validate_sync, file_bytes)

    async def extract_structure(self, file_bytes: bytes) -> dict:
        """
        Extract document structure (headers, body, footers) as a list of nodes
        for the generation preview UI.

        Node kinds returned:
          - "paragraph" — plain body text. `spans` carries the inline pieces.
          - "heading"   — `level` is 1-6.
          - "list_bullet" / "list_number" — `level` is the indentation depth
            (1 for top-level, 2+ for nested items).
          - "table"     — `rows` is a list of row dicts; each row has `cells`,
            each cell has `nodes` (a nested list of paragraph/heading/list
            nodes — table-in-table is intentionally NOT recursed in this
            iteration to keep the preview bounded).

        Empty paragraphs are skipped. Order between paragraphs and tables is
        preserved by walking the underlying XML (CT_P / CT_Tbl) in document
        order.
        """
        # Local imports — these are private-ish python-docx internals but the
        # block-iteration recipe is the standard pattern in the docx ecosystem.
        from docx.oxml.ns import qn
        from docx.table import Table as _DocxTable
        from docx.text.paragraph import Paragraph as _DocxParagraph

        # Same regex used by extract_variables / validate so the three stay in sync.
        # Note: capture the WHOLE placeholder ({{ name }}) plus the inner name.
        placeholder_re = re.compile(r"(\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\})")

        def _spans_for(text: str) -> list[dict]:
            """Split paragraph text into a list of plain-text and placeholder spans."""
            spans: list[dict] = []
            last_end = 0
            for m in placeholder_re.finditer(text):
                if m.start() > last_end:
                    spans.append({"text": text[last_end : m.start()], "variable": None})
                spans.append({"text": m.group(1), "variable": m.group(2)})
                last_end = m.end()
            if last_end < len(text):
                spans.append({"text": text[last_end:], "variable": None})
            return spans

        def _classify(para) -> tuple[str, int]:
            """Detect heading / list / paragraph via the paragraph style name."""
            style_name = ""
            try:
                if para.style is not None and para.style.name:
                    style_name = para.style.name
            except AttributeError:
                # Defensive: some malformed docx files don't expose style cleanly
                style_name = ""

            if style_name.startswith("Heading"):
                # "Heading 1" → level 1, "Heading 2" → 2, ...
                # If no number present (just "Heading"), default to level 1.
                try:
                    level = int(style_name.split()[-1])
                    if level < 1 or level > 6:
                        level = 1
                except (ValueError, IndexError):
                    level = 1
                return ("heading", level)
            if style_name == "Title":
                return ("heading", 1)
            if style_name.startswith("List Bullet"):
                # "List Bullet" → 1, "List Bullet 2" → 2, ...
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = 1
                return ("list_bullet", level)
            if style_name.startswith("List Number"):
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = 1
                return ("list_number", level)
            return ("paragraph", 0)

        def _node_for_para(para) -> dict | None:
            """Build a structure node from a python-docx paragraph, or None if empty."""
            text = para.text
            if not text or not text.strip():
                return None
            kind, level = _classify(para)
            return {"kind": kind, "level": level, "spans": _spans_for(text)}

        def _node_for_table(table) -> dict | None:
            """Build a table node. Cells contain nested paragraph/heading/list
            nodes; nested tables inside cells are intentionally skipped to
            avoid unbounded recursion in pathological documents."""
            rows: list[dict] = []
            any_content = False
            for row in table.rows:
                cells: list[dict] = []
                for cell in row.cells:
                    cell_nodes: list[dict] = []
                    for para in cell.paragraphs:
                        node = _node_for_para(para)
                        if node is not None:
                            cell_nodes.append(node)
                            any_content = True
                    cells.append({"nodes": cell_nodes})
                rows.append({"cells": cells})
            if not any_content:
                # Skip tables whose every cell is blank — they are usually
                # spacing artifacts in legal templates.
                return None
            return {"kind": "table", "level": 0, "spans": [], "rows": rows}

        # Body container lives at `doc.element.body`; headers/footers expose
        # their container directly at `_element`. The block parent passed to
        # Paragraph/Table must be the *python-docx parent object* (Document,
        # _Header, _Footer) — not None — so that `paragraph.style` resolves
        # via `parent.part.styles`.

        p_tag = qn("w:p")
        tbl_tag = qn("w:tbl")

        def _walk_blocks(parent_obj, container_elm) -> list[dict]:
            """Iterate `container_elm` children in document order and emit
            structure nodes. `parent_obj` is the python-docx owner used for
            style resolution."""
            out: list[dict] = []
            for child in container_elm.iterchildren():
                if child.tag == p_tag:
                    node = _node_for_para(_DocxParagraph(child, parent_obj))
                elif child.tag == tbl_tag:
                    node = _node_for_table(_DocxTable(child, parent_obj))
                else:
                    node = None
                if node is not None:
                    out.append(node)
            return out

        def _extract(data: bytes) -> dict:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp.write(data)
                tmp_path = tmp.name
            try:
                doc = Document(tmp_path)

                body_nodes = _walk_blocks(doc, doc.element.body)

                # Sections can have distinct headers/footers; we flatten them
                # into a single list each. The first iteration of the preview
                # treats the document as one logical scope.
                header_nodes: list[dict] = []
                footer_nodes: list[dict] = []
                for section in doc.sections:
                    if section.header:
                        header_nodes.extend(
                            _walk_blocks(section.header, section.header._element)
                        )
                    if section.footer:
                        footer_nodes.extend(
                            _walk_blocks(section.footer, section.footer._element)
                        )

                return {
                    "headers": header_nodes,
                    "body": body_nodes,
                    "footers": footer_nodes,
                }
            finally:
                os.unlink(tmp_path)

        return await asyncio.to_thread(_extract, file_bytes)

    async def auto_fix(self, file_bytes: bytes) -> bytes:
        """
        Auto-fix fixable issues in a template file.

        Fixes: no-space variables, uppercase, camelCase -> snake_case.
        Does NOT fix: special chars, split formatting.
        Returns the corrected file bytes.

        IMPORTANT: Works at the paragraph level (not run level) to handle variables
        that Word splits across multiple runs due to mixed formatting. After applying
        fixes to the full paragraph text, the runs are rebuilt: all run texts are
        cleared and the entire fixed text is placed in the first run (preserving
        that run's font properties).
        """

        def _fix_text(text: str) -> str:
            """Apply all fixable substitutions to a full paragraph text string."""
            # Fix {{var}} -> {{ snake_case }}
            # NOTE: pass the raw match to _camel_to_snake (it lowercases internally).
            # Calling .lower() BEFORE _camel_to_snake destroys case info needed for
            # splitting — e.g. 'UserName'.lower() = 'username' → never splits.
            text = re.sub(
                r"\{\{(\w+)\}\}",
                lambda m: "{{ " + _camel_to_snake(m.group(1)) + " }}",
                text,
            )
            # Fix {{ Var }} or {{ camelCase }} -> {{ snake_case }}
            text = re.sub(
                r"\{\{\s+(\w+)\s+\}\}",
                lambda m: "{{ " + _camel_to_snake(m.group(1)) + " }}",
                text,
            )
            return text

        def _fix_paragraph(para) -> None:
            """
            Fix all variable names in a paragraph, working at the paragraph level
            to correctly handle variables split across multiple runs by Word.

            Strategy:
            1. Concatenate all run texts to get the full paragraph text.
            2. Apply regex fixes to the full text.
            3. If the text changed, rebuild runs: clear all runs, put the fixed
               full text in the first run (keeping its font properties intact).
            """
            if not para.runs:
                return

            full_text = para.text  # concatenates all runs
            fixed_text = _fix_text(full_text)

            if fixed_text == full_text:
                return  # nothing to fix

            # Save the first run's font properties before modifying anything
            first_run = para.runs[0]
            bold = first_run.bold
            italic = first_run.italic
            underline = first_run.underline
            font_name = first_run.font.name
            font_size = first_run.font.size
            font_color = first_run.font.color.rgb if first_run.font.color and first_run.font.color.type else None

            # Clear ALL runs so none of the old fragmented text remains
            for run in para.runs:
                run.text = ""

            # Place the entire fixed text in the first run and restore its font
            first_run.text = fixed_text
            first_run.bold = bold
            first_run.italic = italic
            first_run.underline = underline
            if font_name:
                first_run.font.name = font_name
            if font_size:
                first_run.font.size = font_size
            if font_color is not None:
                from docx.shared import RGBColor
                first_run.font.color.rgb = font_color

        def _auto_fix_sync(data: bytes) -> bytes:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp.write(data)
                tmp_path = tmp.name

            try:
                doc = Document(tmp_path)

                # Collect all paragraphs: body, headers, footers, AND table cells
                all_paragraphs = list(doc.paragraphs)
                for section in doc.sections:
                    if section.header:
                        all_paragraphs.extend(section.header.paragraphs)
                    if section.footer:
                        all_paragraphs.extend(section.footer.paragraphs)

                # Tables: body tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            all_paragraphs.extend(cell.paragraphs)

                # Tables inside headers and footers
                for section in doc.sections:
                    if section.header:
                        for table in section.header.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    all_paragraphs.extend(cell.paragraphs)
                    if section.footer:
                        for table in section.footer.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    all_paragraphs.extend(cell.paragraphs)

                for para in all_paragraphs:
                    _fix_paragraph(para)

                output = io.BytesIO()
                doc.save(output)
                output.seek(0)
                return output.read()
            finally:
                os.unlink(tmp_path)

        return await asyncio.to_thread(_auto_fix_sync, file_bytes)
