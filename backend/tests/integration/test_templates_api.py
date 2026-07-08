"""Integration tests — /api/v1/templates/* endpoints.

All tests use the fakes wired up in the integration conftest.  The template
service override provides FakeTemplateRepository + FakeStorageService +
FakeTemplateEngine.

Template upload also calls get_template_engine() directly (outside DI) for
validation.  We monkeypatch that call to return our FakeTemplateEngine so
validation always passes.
"""

from __future__ import annotations

import io
import uuid

import pytest

from app.domain.entities import User
from app.presentation.middleware.tenant import CurrentUser, get_current_user
from tests.fakes import FakeTemplateRepository, FakeUserRepository

# Stable IDs from integration conftest
CONFTEST_USER_ID = uuid.UUID("aaaaaaaa-aaaa-aaaa-aaaa-aaaaaaaaaaaa")
CONFTEST_TENANT_ID = uuid.UUID("bbbbbbbb-bbbb-bbbb-bbbb-bbbbbbbbbbbb")

# A second user in the same tenant (different from the admin test_user)
USER_B_ID = uuid.UUID("11111111-1111-1111-1111-111111111111")


# ── Unauthenticated access ────────────────────────────────────────────────────


@pytest.mark.asyncio
async def test_list_templates_without_auth_returns_401(async_client, app):
    from app.presentation.middleware.tenant import get_current_user

    original = app.dependency_overrides.pop(get_current_user, None)
    try:
        response = await async_client.get("/api/v1/templates")
        assert response.status_code == 401
    finally:
        if original is not None:
            app.dependency_overrides[get_current_user] = original


@pytest.mark.asyncio
async def test_upload_template_without_auth_returns_401(async_client, app):
    from app.presentation.middleware.tenant import get_current_user

    original = app.dependency_overrides.pop(get_current_user, None)
    try:
        response = await async_client.post(
            "/api/v1/templates/upload",
            data={"name": "Test"},
            files={"file": ("test.docx", b"fake-bytes", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert response.status_code == 401
    finally:
        if original is not None:
            app.dependency_overrides[get_current_user] = original


# ── Authenticated list ────────────────────────────────────────────────────────


@pytest.mark.asyncio
async def test_list_templates_returns_200(async_client, auth_headers):
    response = await async_client.get("/api/v1/templates", headers=auth_headers)
    assert response.status_code == 200
    data = response.json()
    assert "items" in data
    assert "total" in data
    assert isinstance(data["items"], list)


@pytest.mark.asyncio
async def test_list_templates_pagination_fields_present(async_client, auth_headers):
    response = await async_client.get(
        "/api/v1/templates", headers=auth_headers, params={"page": 1, "size": 10}
    )
    assert response.status_code == 200
    data = response.json()
    assert data["page"] == 1
    assert data["size"] == 10


# ── Template upload ───────────────────────────────────────────────────────────


@pytest.mark.asyncio
async def test_upload_template_creates_template(
    async_client, auth_headers, monkeypatch, fake_template_engine
):
    """Uploading a template returns 201 and the template appears in the list."""
    # Configure the shared fake engine (used by both the service and the direct call)
    fake_template_engine.variables_to_return = ["name", "date"]

    # Patch get_template_engine() used directly inside the upload endpoint for validation
    monkeypatch.setattr(
        "app.presentation.api.v1.templates.get_template_engine",
        lambda: fake_template_engine,
    )

    docx_bytes = b"fake-docx-content"
    response = await async_client.post(
        "/api/v1/templates/upload",
        headers=auth_headers,
        data={"name": "My Integration Template"},
        files={
            "file": (
                "template.docx",
                io.BytesIO(docx_bytes),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 201
    data = response.json()
    assert data["name"] == "My Integration Template"
    assert "id" in data
    assert data["version"] == 1
    assert set(data["variables"]) == {"name", "date"}

    # Reset for subsequent tests
    fake_template_engine.variables_to_return = []


@pytest.mark.asyncio
async def test_upload_template_appears_in_list(
    async_client, auth_headers, monkeypatch, fake_template_engine
):
    """After upload, the template is returned in GET /templates."""
    fake_template_engine.variables_to_return = ["title"]
    monkeypatch.setattr(
        "app.presentation.api.v1.templates.get_template_engine",
        lambda: fake_template_engine,
    )

    docx_bytes = b"fake-docx-for-list-test"
    upload_response = await async_client.post(
        "/api/v1/templates/upload",
        headers=auth_headers,
        data={"name": "ListableTemplate"},
        files={
            "file": (
                "template.docx",
                io.BytesIO(docx_bytes),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert upload_response.status_code == 201
    created_id = upload_response.json()["id"]

    # Use the unique name as a search filter so the assertion is not affected by
    # the number of templates accumulated in the session-scoped fake repository.
    # Without this, a full test suite run can push the new template past page 1
    # (default size=20), causing a false negative while the test passes in isolation.
    list_response = await async_client.get(
        "/api/v1/templates",
        headers=auth_headers,
        params={"search": "ListableTemplate"},
    )
    assert list_response.status_code == 200
    ids = [item["id"] for item in list_response.json()["items"]]
    assert created_id in ids

    # Reset
    fake_template_engine.variables_to_return = []


@pytest.mark.asyncio
async def test_upload_empty_file_returns_400(async_client, auth_headers, monkeypatch, fake_template_engine):
    monkeypatch.setattr(
        "app.presentation.api.v1.templates.get_template_engine",
        lambda: fake_template_engine,
    )

    response = await async_client.post(
        "/api/v1/templates/upload",
        headers=auth_headers,
        data={"name": "EmptyTemplate"},
        files={
            "file": (
                "empty.docx",
                io.BytesIO(b""),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 400


# ── Task 5.5: Template visibility / sharing integration tests ─────────────────


def _upload_template_as_user(
    template_name: str,
    owner_id: uuid.UUID,
    fake_template_repo: FakeTemplateRepository,
) -> uuid.UUID:
    """Directly seed a template in the fake repo for a specific owner.

    Returns the template_id.
    """
    from datetime import datetime, timezone
    from app.domain.entities import Template, TemplateVersion

    template_id = uuid.uuid4()
    version_id = uuid.uuid4()
    now = datetime.now(timezone.utc)

    version = TemplateVersion(
        id=version_id,
        tenant_id=CONFTEST_TENANT_ID,
        template_id=template_id,
        version=1,
        minio_path=f"{CONFTEST_TENANT_ID}/{template_id}/v1/template.docx",
        variables=["name"],
        created_at=now,
    )
    template = Template(
        id=template_id,
        tenant_id=CONFTEST_TENANT_ID,
        name=template_name,
        description=None,
        current_version=1,
        created_by=owner_id,
        versions=[version],
        created_at=now,
        updated_at=now,
    )
    fake_template_repo._templates[template_id] = template
    fake_template_repo._versions[version_id] = version
    return template_id


@pytest.mark.asyncio
async def test_private_by_default_user_b_cannot_list(
    async_client, app, auth_headers, fake_template_repo
):
    """User B cannot see user A's template in the list (private by default)."""
    # Seed a template owned by the conftest admin user (user A)
    tpl_id = _upload_template_as_user(
        "PrivateByDefault", CONFTEST_USER_ID, fake_template_repo
    )

    # Override current_user to user B (a regular user)
    user_b = CurrentUser(
        user_id=USER_B_ID,
        tenant_id=CONFTEST_TENANT_ID,
        role="document_generator",
    )

    async def override_as_user_b():
        return user_b

    original = app.dependency_overrides.get(get_current_user)
    app.dependency_overrides[get_current_user] = override_as_user_b
    try:
        response = await async_client.get("/api/v1/templates", headers=auth_headers)
        assert response.status_code == 200
        ids = [item["id"] for item in response.json()["items"]]
        assert str(tpl_id) not in ids
    finally:
        if original is not None:
            app.dependency_overrides[get_current_user] = original
        else:
            app.dependency_overrides.pop(get_current_user, None)


@pytest.mark.asyncio
async def test_private_by_default_user_b_gets_403_on_get(
    async_client, app, auth_headers, fake_template_repo
):
    """User B gets 403 when trying to GET user A's private template."""
    tpl_id = _upload_template_as_user(
        "PrivateGet403", CONFTEST_USER_ID, fake_template_repo
    )

    user_b = CurrentUser(
        user_id=USER_B_ID,
        tenant_id=CONFTEST_TENANT_ID,
        role="document_generator",
    )

    async def override_as_user_b():
        return user_b

    original = app.dependency_overrides.get(get_current_user)
    app.dependency_overrides[get_current_user] = override_as_user_b
    try:
        response = await async_client.get(
            f"/api/v1/templates/{tpl_id}", headers=auth_headers
        )
        assert response.status_code == 403
    finally:
        if original is not None:
            app.dependency_overrides[get_current_user] = original
        else:
            app.dependency_overrides.pop(get_current_user, None)


@pytest.mark.asyncio
async def test_after_sharing_user_b_sees_in_list(
    async_client, app, auth_headers, fake_template_repo
):
    """After sharing, user B sees the template in the list."""
    tpl_id = _upload_template_as_user(
        "SharedVisible", CONFTEST_USER_ID, fake_template_repo
    )

    # Grant share to user B directly in fake repo
    await fake_template_repo.add_share(
        template_id=tpl_id,
        user_id=USER_B_ID,
        tenant_id=CONFTEST_TENANT_ID,
        shared_by=CONFTEST_USER_ID,
    )

    user_b = CurrentUser(
        user_id=USER_B_ID,
        tenant_id=CONFTEST_TENANT_ID,
        role="document_generator",
    )

    async def override_as_user_b():
        return user_b

    original = app.dependency_overrides.get(get_current_user)
    app.dependency_overrides[get_current_user] = override_as_user_b
    try:
        response = await async_client.get("/api/v1/templates", headers=auth_headers)
        assert response.status_code == 200
        ids = [item["id"] for item in response.json()["items"]]
        assert str(tpl_id) in ids
    finally:
        if original is not None:
            app.dependency_overrides[get_current_user] = original
        else:
            app.dependency_overrides.pop(get_current_user, None)


@pytest.mark.asyncio
async def test_shared_user_cannot_upload_new_version(
    async_client, app, auth_headers, fake_template_repo, monkeypatch, fake_template_engine
):
    """User B cannot upload a new version to a template they only have shared access to."""
    tpl_id = _upload_template_as_user(
        "SharedNoVersion", CONFTEST_USER_ID, fake_template_repo
    )
    await fake_template_repo.add_share(
        template_id=tpl_id,
        user_id=USER_B_ID,
        tenant_id=CONFTEST_TENANT_ID,
        shared_by=CONFTEST_USER_ID,
    )

    monkeypatch.setattr(
        "app.presentation.api.v1.templates.get_template_engine",
        lambda: fake_template_engine,
    )
    fake_template_engine.variables_to_return = ["name"]

    user_b = CurrentUser(
        user_id=USER_B_ID,
        tenant_id=CONFTEST_TENANT_ID,
        role="document_generator",
    )

    async def override_as_user_b():
        return user_b

    original = app.dependency_overrides.get(get_current_user)
    app.dependency_overrides[get_current_user] = override_as_user_b
    try:
        response = await async_client.post(
            f"/api/v1/templates/{tpl_id}/versions",
            headers=auth_headers,
            files={
                "file": (
                    "template.docx",
                    io.BytesIO(b"fake-bytes"),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert response.status_code == 403
    finally:
        if original is not None:
            app.dependency_overrides[get_current_user] = original
        else:
            app.dependency_overrides.pop(get_current_user, None)


@pytest.mark.asyncio
async def test_shared_user_cannot_delete_template(
    async_client, app, auth_headers, fake_template_repo
):
    """User B cannot delete a template they only have shared access to."""
    tpl_id = _upload_template_as_user(
        "SharedNoDelete", CONFTEST_USER_ID, fake_template_repo
    )
    await fake_template_repo.add_share(
        template_id=tpl_id,
        user_id=USER_B_ID,
        tenant_id=CONFTEST_TENANT_ID,
        shared_by=CONFTEST_USER_ID,
    )

    user_b = CurrentUser(
        user_id=USER_B_ID,
        tenant_id=CONFTEST_TENANT_ID,
        role="document_generator",
    )

    async def override_as_user_b():
        return user_b

    original = app.dependency_overrides.get(get_current_user)
    app.dependency_overrides[get_current_user] = override_as_user_b
    try:
        response = await async_client.delete(
            f"/api/v1/templates/{tpl_id}", headers=auth_headers
        )
        assert response.status_code == 403
    finally:
        if original is not None:
            app.dependency_overrides[get_current_user] = original
        else:
            app.dependency_overrides.pop(get_current_user, None)


@pytest.mark.asyncio
async def test_list_response_has_access_type_field(
    async_client, auth_headers, fake_template_repo
):
    """Template listing includes access_type field for each item."""
    # Admin user (conftest user) uploads — so access_type should be 'owned'
    _upload_template_as_user(
        "WithAccessType", CONFTEST_USER_ID, fake_template_repo
    )

    response = await async_client.get("/api/v1/templates", headers=auth_headers)
    assert response.status_code == 200
    items = response.json()["items"]
    assert len(items) > 0
    for item in items:
        assert "access_type" in item
        assert item["access_type"] in ("owned", "shared", "admin")


@pytest.mark.asyncio
async def test_upload_response_returns_full_paragraph_contexts_for_multi_var_paragraphs(
    async_client, app, auth_headers, fake_template_repo, fake_storage, fake_audit_repo
):
    """POSTing a real .docx with multi-variable paragraphs must produce variables_meta
    where co-occurring variables share the IDENTICAL full paragraph context string,
    with no truncation markers.  This locks the handler→service→engine→schema wiring."""
    import io as _io

    from docx import Document as _Document

    from app.application.services.audit_service import AuditService
    from app.application.services.template_service import TemplateService
    from app.application.services import get_template_service
    from app.infrastructure.templating.docxtpl_engine import DocxTemplateEngine

    # Build a real .docx in memory with two paragraphs:
    #  - Para 1: three variables in a single long sentence (tests shared-context invariant)
    #  - Para 2: one isolated variable (ensures paragraphs are not conflated)
    multi_var_para = (
        "La empresa {{ nombre_empresa }} con NIT {{ numero }}, "
        "en la ciudad de {{ ciudad }}."
    )
    solo_var_para = "Fecha de vigencia del contrato: {{ fecha_vigencia }}."

    doc = _Document()
    doc.add_paragraph().add_run(multi_var_para)
    doc.add_paragraph().add_run(solo_var_para)
    buf = _io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    docx_bytes = buf.read()

    # Wire the real engine through both the validation call and the service
    real_engine = DocxTemplateEngine()

    async def _override_template_service() -> TemplateService:
        _audit_svc = AuditService(audit_repo=fake_audit_repo)
        return TemplateService(
            repository=fake_template_repo,
            storage=fake_storage,
            engine=real_engine,
            audit_service=_audit_svc,
        )

    original_override = app.dependency_overrides.get(get_template_service)
    app.dependency_overrides[get_template_service] = _override_template_service

    import app.presentation.api.v1.templates as _tpl_module
    original_get_engine = _tpl_module.get_template_engine
    _tpl_module.get_template_engine = lambda: real_engine

    try:
        response = await async_client.post(
            "/api/v1/templates/upload",
            headers=auth_headers,
            data={"name": "Contrato multi-variable"},
            files={
                "file": (
                    "contrato.docx",
                    _io.BytesIO(docx_bytes),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

        assert response.status_code == 201, f"Upload failed: {response.text}"
        template_id = response.json()["id"]

        # Retrieve the full template detail (TemplateUploadResponse lacks versions)
        detail_response = await async_client.get(
            f"/api/v1/templates/{template_id}", headers=auth_headers
        )
        assert detail_response.status_code == 200
        data = detail_response.json()

        vm = data["versions"][0]["variables_meta"]
        assert isinstance(vm, list) and len(vm) > 0, "variables_meta must be non-empty"

        nombre_meta = next(m for m in vm if m["name"] == "nombre_empresa")
        numero_meta  = next(m for m in vm if m["name"] == "numero")
        ciudad_meta  = next(m for m in vm if m["name"] == "ciudad")

        # All three co-occurring variables must share the exact same context string
        assert (
            nombre_meta["contexts"][0]
            == numero_meta["contexts"][0]
            == ciudad_meta["contexts"][0]
        ), "Co-occurring variables must share the identical full paragraph context"

        # No truncation markers in any context
        for var_meta in [nombre_meta, numero_meta, ciudad_meta]:
            for ctx in var_meta["contexts"]:
                assert "..." not in ctx, (
                    f"Context for '{var_meta['name']}' must not be truncated, got: {ctx!r}"
                )

        # Isolated variable in second paragraph must NOT share context with multi-var para
        fecha_meta = next(m for m in vm if m["name"] == "fecha_vigencia")
        assert fecha_meta["contexts"][0] != nombre_meta["contexts"][0], (
            "Variables from different paragraphs must not share the same context string"
        )
    finally:
        if original_override is not None:
            app.dependency_overrides[get_template_service] = original_override
        else:
            app.dependency_overrides.pop(get_template_service, None)
        _tpl_module.get_template_engine = original_get_engine


@pytest.mark.asyncio
async def test_admin_can_delete_any_template(
    async_client, app, auth_headers, fake_template_repo
):
    """Admin can delete a template owned by another user, and it disappears from the list."""
    # Seed a template owned by USER_B (not the admin test user)
    tpl_id = _upload_template_as_user(
        "AdminDeleteTarget", USER_B_ID, fake_template_repo
    )

    # The default test user (CONFTEST_USER_ID) has role "admin" per conftest
    response = await async_client.delete(
        f"/api/v1/templates/{tpl_id}", headers=auth_headers
    )
    assert response.status_code in (200, 204)

    # Verify the template is gone from the list (admin sees all)
    list_response = await async_client.get("/api/v1/templates", headers=auth_headers)
    assert list_response.status_code == 200
    ids = [item["id"] for item in list_response.json()["items"]]
    assert str(tpl_id) not in ids


# ── shared_by_email in GET /templates/{id} ────────────────────────────────────

# A dedicated owner for these tests to avoid collisions with the admin test user
OWNER_FOR_SHARE_EMAIL_TESTS_ID = uuid.UUID("77777777-7777-7777-7777-777777777777")
SHARED_RECIPIENT_ID = uuid.UUID("88888888-8888-8888-8888-888888888888")


async def _ensure_user(
    fake_user_repo: FakeUserRepository,
    user_id: uuid.UUID,
    email: str,
    role: str = "user",
) -> User:
    """Upsert a user in the fake repo. Returns the user."""
    existing = await fake_user_repo.get_by_id(user_id)
    if existing is not None:
        return existing
    user = User(
        id=user_id,
        tenant_id=CONFTEST_TENANT_ID,
        email=email,
        hashed_password="hashed",
        full_name="Test User",
        role=role,
        is_active=True,
    )
    await fake_user_repo.create(user)
    return user


@pytest.mark.asyncio
async def test_get_template_includes_shared_by_email_when_access_type_is_shared(
    async_client,
    app,
    auth_headers,
    fake_template_repo: FakeTemplateRepository,
    fake_user_repo: FakeUserRepository,
):
    """GET /templates/{id} as the share recipient returns shared_by_email == owner's email."""
    # Seed owner (User A) and recipient (User B) in the user repo
    owner = await _ensure_user(
        fake_user_repo,
        OWNER_FOR_SHARE_EMAIL_TESTS_ID,
        "owner-shared-by@example.com",
        role="user",
    )
    recipient = await _ensure_user(
        fake_user_repo,
        SHARED_RECIPIENT_ID,
        "recipient-shared@example.com",
        role="user",
    )

    # Seed a template owned by User A
    tpl_id = _upload_template_as_user("SharedEmailTemplate", owner.id, fake_template_repo)

    # Share it with User B (recipient)
    await fake_template_repo.add_share(
        template_id=tpl_id,
        user_id=recipient.id,
        tenant_id=CONFTEST_TENANT_ID,
        shared_by=owner.id,
    )

    # Override current user → User B (recipient) calls GET /templates/{id}
    recipient_current_user = CurrentUser(
        user_id=recipient.id,
        tenant_id=CONFTEST_TENANT_ID,
        role="user",
    )

    async def override_as_recipient():
        return recipient_current_user

    original = app.dependency_overrides.get(get_current_user)
    app.dependency_overrides[get_current_user] = override_as_recipient
    try:
        response = await async_client.get(
            f"/api/v1/templates/{tpl_id}", headers=auth_headers
        )
        assert response.status_code == 200
        data = response.json()
        assert data["access_type"] == "shared"
        assert data["shared_by_email"] == owner.email
    finally:
        if original is not None:
            app.dependency_overrides[get_current_user] = original
        else:
            app.dependency_overrides.pop(get_current_user, None)


@pytest.mark.asyncio
async def test_get_template_shared_by_email_is_none_for_owner(
    async_client,
    auth_headers,
    fake_template_repo: FakeTemplateRepository,
    fake_user_repo: FakeUserRepository,
):
    """GET /templates/{id} as the owner returns shared_by_email == null."""
    # The default test user (CONFTEST_USER_ID, admin) owns this template
    tpl_id = _upload_template_as_user(
        "OwnedNoSharedByEmail", CONFTEST_USER_ID, fake_template_repo
    )

    response = await async_client.get(
        f"/api/v1/templates/{tpl_id}", headers=auth_headers
    )
    assert response.status_code == 200
    data = response.json()
    assert data["access_type"] in ("owned", "admin")
    assert data["shared_by_email"] is None


# ── variable types PATCH endpoint ────────────────────────────────────────────


def _seed_template_with_variables(
    variables: list[str],
    owner_id: uuid.UUID,
    fake_template_repo: FakeTemplateRepository,
    template_name: str = "VarTypeTemplate",
) -> tuple[uuid.UUID, uuid.UUID]:
    """Seed a template with the given variable names. Returns (template_id, version_id)."""
    from datetime import datetime, timezone
    from app.domain.entities import Template, TemplateVersion

    template_id = uuid.uuid4()
    version_id = uuid.uuid4()
    now = datetime.now(timezone.utc)

    version = TemplateVersion(
        id=version_id,
        tenant_id=CONFTEST_TENANT_ID,
        template_id=template_id,
        version=1,
        minio_path=f"{CONFTEST_TENANT_ID}/{template_id}/v1/template.docx",
        variables=variables,
        variables_meta=[{"name": v, "contexts": [f"context for {v}"]} for v in variables],
        created_at=now,
    )
    template = Template(
        id=template_id,
        tenant_id=CONFTEST_TENANT_ID,
        name=template_name,
        description=None,
        current_version=1,
        created_by=owner_id,
        versions=[version],
        created_at=now,
        updated_at=now,
    )
    fake_template_repo._templates[template_id] = template
    fake_template_repo._versions[version_id] = version
    return template_id, version_id


@pytest.mark.asyncio
async def test_owner_can_update_variable_types(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    """Owner can PATCH variable types and get the updated version back."""
    tpl_id, ver_id = _seed_template_with_variables(
        ["monto_total", "moneda"], CONFTEST_USER_ID, fake_template_repo, "VarTypeOwner"
    )

    response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}/versions/{ver_id}/variables-meta",
        headers=auth_headers,
        json={"overrides": [{"name": "monto_total", "type": "decimal"}]},
    )
    assert response.status_code == 200
    data = response.json()
    meta_by_name = {m["name"]: m for m in data["variables_meta"]}
    assert meta_by_name["monto_total"]["type"] == "decimal"
    # moneda was not in overrides — should default to "text"
    assert meta_by_name["moneda"]["type"] == "text"


@pytest.mark.asyncio
async def test_admin_cannot_update_types_of_template_they_do_not_own(
    async_client, app, auth_headers, fake_template_repo: FakeTemplateRepository
):
    """Admin calling PATCH on someone else's template gets 403."""
    other_owner = uuid.UUID("cccccccc-cccc-cccc-cccc-cccccccccccc")
    tpl_id, ver_id = _seed_template_with_variables(
        ["campo"], other_owner, fake_template_repo, "VarTypeAdminDenied"
    )

    # Conftest user is admin but NOT the owner — must get 403
    response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}/versions/{ver_id}/variables-meta",
        headers=auth_headers,
        json={"overrides": [{"name": "campo", "type": "integer"}]},
    )
    assert response.status_code == 403


@pytest.mark.asyncio
async def test_recipient_cannot_update_types(
    async_client, app, auth_headers, fake_template_repo: FakeTemplateRepository
):
    """A user with only shared access gets 403 when trying to PATCH variable types."""
    tpl_id, ver_id = _seed_template_with_variables(
        ["campo"], CONFTEST_USER_ID, fake_template_repo, "VarTypeRecipientDenied"
    )
    await fake_template_repo.add_share(
        template_id=tpl_id,
        user_id=USER_B_ID,
        tenant_id=CONFTEST_TENANT_ID,
        shared_by=CONFTEST_USER_ID,
    )

    user_b = CurrentUser(
        user_id=USER_B_ID,
        tenant_id=CONFTEST_TENANT_ID,
        role="document_generator",
    )

    async def override_as_user_b():
        return user_b

    original = app.dependency_overrides.get(get_current_user)
    app.dependency_overrides[get_current_user] = override_as_user_b
    try:
        response = await async_client.patch(
            f"/api/v1/templates/{tpl_id}/versions/{ver_id}/variables-meta",
            headers=auth_headers,
            json={"overrides": [{"name": "campo", "type": "integer"}]},
        )
        assert response.status_code == 403
    finally:
        if original is not None:
            app.dependency_overrides[get_current_user] = original
        else:
            app.dependency_overrides.pop(get_current_user, None)


@pytest.mark.asyncio
async def test_select_without_options_returns_422(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    """PATCH with type=select but no options returns 422."""
    tpl_id, ver_id = _seed_template_with_variables(
        ["moneda"], CONFTEST_USER_ID, fake_template_repo, "VarTypeSelectNoOpts"
    )

    response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}/versions/{ver_id}/variables-meta",
        headers=auth_headers,
        json={"overrides": [{"name": "moneda", "type": "select"}]},
    )
    assert response.status_code == 422


@pytest.mark.asyncio
async def test_invalid_type_returns_422(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    """PATCH with an invalid type value returns 422."""
    tpl_id, ver_id = _seed_template_with_variables(
        ["campo"], CONFTEST_USER_ID, fake_template_repo, "VarTypeInvalidType"
    )

    response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}/versions/{ver_id}/variables-meta",
        headers=auth_headers,
        json={"overrides": [{"name": "campo", "type": "boolean"}]},
    )
    assert response.status_code == 422


@pytest.mark.asyncio
async def test_unknown_variable_name_in_overrides_is_ignored(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    """Overrides for unknown variable names are silently ignored (no error)."""
    tpl_id, ver_id = _seed_template_with_variables(
        ["conocido"], CONFTEST_USER_ID, fake_template_repo, "VarTypeUnknownIgnored"
    )

    response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}/versions/{ver_id}/variables-meta",
        headers=auth_headers,
        json={"overrides": [{"name": "no_existe", "type": "integer"}]},
    )
    assert response.status_code == 200
    data = response.json()
    meta_by_name = {m["name"]: m for m in data["variables_meta"]}
    assert "no_existe" not in meta_by_name
    assert "conocido" in meta_by_name


@pytest.mark.asyncio
async def test_existing_meta_without_type_field_defaults_to_text(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    """Legacy variables_meta rows without a 'type' field serialize as type='text'."""
    tpl_id, ver_id = _seed_template_with_variables(
        ["legacy_var"], CONFTEST_USER_ID, fake_template_repo, "VarTypeLegacy"
    )

    # The seeded meta has no 'type' key — call PATCH with empty overrides
    response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}/versions/{ver_id}/variables-meta",
        headers=auth_headers,
        json={"overrides": []},
    )
    assert response.status_code == 200
    data = response.json()
    meta_by_name = {m["name"]: m for m in data["variables_meta"]}
    assert meta_by_name["legacy_var"]["type"] == "text"


@pytest.mark.asyncio
async def test_partial_override_preserves_other_variables(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    """Only overridden variables change; others keep their existing data."""
    tpl_id, ver_id = _seed_template_with_variables(
        ["nombre", "monto", "fecha"], CONFTEST_USER_ID, fake_template_repo, "VarTypePartial"
    )

    response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}/versions/{ver_id}/variables-meta",
        headers=auth_headers,
        json={"overrides": [{"name": "monto", "type": "decimal"}]},
    )
    assert response.status_code == 200
    data = response.json()
    meta_by_name = {m["name"]: m for m in data["variables_meta"]}

    assert meta_by_name["monto"]["type"] == "decimal"
    assert meta_by_name["nombre"]["type"] == "text"
    assert meta_by_name["fecha"]["type"] == "text"


# ── computed variables (formulas + functions) ────────────────────────────────


@pytest.mark.asyncio
async def test_owner_can_save_formula_computed_variable(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    """Owner can PATCH a formula computed variable; type is normalized to decimal."""
    tpl_id, ver_id = _seed_template_with_variables(
        ["monto", "total_con_iva"], CONFTEST_USER_ID, fake_template_repo, "ComputedFormulaOwner"
    )

    response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}/versions/{ver_id}/variables-meta",
        headers=auth_headers,
        json={
            "overrides": [
                {"name": "monto", "type": "decimal"},
                {
                    "name": "total_con_iva",
                    "type": "text",  # deliberately wrong — server normalizes to decimal
                    "computed": {
                        "kind": "formula",
                        "source": "monto",
                        "operator": "+",
                        "operand": 100,
                    },
                },
            ]
        },
    )
    assert response.status_code == 200
    data = response.json()
    meta_by_name = {m["name"]: m for m in data["variables_meta"]}
    assert meta_by_name["total_con_iva"]["type"] == "decimal"
    assert meta_by_name["total_con_iva"]["computed"]["kind"] == "formula"


@pytest.mark.asyncio
async def test_owner_can_save_function_computed_variable(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    """Owner can PATCH a number_to_words computed variable; type normalizes to text."""
    tpl_id, ver_id = _seed_template_with_variables(
        ["monto", "monto_en_letras"],
        CONFTEST_USER_ID,
        fake_template_repo,
        "ComputedFunctionOwner",
    )

    response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}/versions/{ver_id}/variables-meta",
        headers=auth_headers,
        json={
            "overrides": [
                {"name": "monto", "type": "decimal"},
                {
                    "name": "monto_en_letras",
                    "type": "integer",  # deliberately wrong — server normalizes to text
                    "computed": {
                        "kind": "function",
                        "function": "number_to_words",
                        "source": "monto",
                    },
                },
            ]
        },
    )
    assert response.status_code == 200
    data = response.json()
    meta_by_name = {m["name"]: m for m in data["variables_meta"]}
    assert meta_by_name["monto_en_letras"]["type"] == "text"
    assert meta_by_name["monto_en_letras"]["computed"]["function"] == "number_to_words"


@pytest.mark.asyncio
async def test_computed_with_unknown_source_returns_422(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    tpl_id, ver_id = _seed_template_with_variables(
        ["total"], CONFTEST_USER_ID, fake_template_repo, "ComputedUnknownSource"
    )

    response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}/versions/{ver_id}/variables-meta",
        headers=auth_headers,
        json={
            "overrides": [
                {
                    "name": "total",
                    "type": "decimal",
                    "computed": {
                        "kind": "formula",
                        "source": "no_existe",
                        "operator": "+",
                        "operand": 1,
                    },
                }
            ]
        },
    )
    assert response.status_code == 422


@pytest.mark.asyncio
async def test_computed_with_non_numeric_source_returns_422(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    tpl_id, ver_id = _seed_template_with_variables(
        ["nombre", "total"], CONFTEST_USER_ID, fake_template_repo, "ComputedNonNumericSource"
    )

    response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}/versions/{ver_id}/variables-meta",
        headers=auth_headers,
        json={
            "overrides": [
                {"name": "nombre", "type": "text"},
                {
                    "name": "total",
                    "type": "decimal",
                    "computed": {
                        "kind": "formula",
                        "source": "nombre",
                        "operator": "+",
                        "operand": 1,
                    },
                },
            ]
        },
    )
    assert response.status_code == 422


@pytest.mark.asyncio
async def test_computed_division_by_zero_operand_returns_422(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    tpl_id, ver_id = _seed_template_with_variables(
        ["monto", "total"], CONFTEST_USER_ID, fake_template_repo, "ComputedDivByZero"
    )

    response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}/versions/{ver_id}/variables-meta",
        headers=auth_headers,
        json={
            "overrides": [
                {
                    "name": "total",
                    "type": "decimal",
                    "computed": {
                        "kind": "formula",
                        "source": "monto",
                        "operator": "/",
                        "operand": 0,
                    },
                }
            ]
        },
    )
    assert response.status_code == 422


@pytest.mark.asyncio
async def test_computed_unknown_function_returns_422(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    tpl_id, ver_id = _seed_template_with_variables(
        ["monto", "total"], CONFTEST_USER_ID, fake_template_repo, "ComputedUnknownFunction"
    )

    response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}/versions/{ver_id}/variables-meta",
        headers=auth_headers,
        json={
            "overrides": [
                {
                    "name": "total",
                    "type": "text",
                    "computed": {
                        "kind": "function",
                        "function": "not_a_real_function",
                        "source": "monto",
                    },
                }
            ]
        },
    )
    assert response.status_code == 422


@pytest.mark.asyncio
async def test_computed_chained_source_returns_422(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    """A computed variable's source must not itself be computed (v1: no chaining)."""
    tpl_id, ver_id = _seed_template_with_variables(
        ["monto", "total", "total_en_letras"],
        CONFTEST_USER_ID,
        fake_template_repo,
        "ComputedChained",
    )

    response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}/versions/{ver_id}/variables-meta",
        headers=auth_headers,
        json={
            "overrides": [
                {"name": "monto", "type": "decimal"},
                {
                    "name": "total",
                    "type": "decimal",
                    "computed": {
                        "kind": "formula",
                        "source": "monto",
                        "operator": "+",
                        "operand": 1,
                    },
                },
                {
                    "name": "total_en_letras",
                    "type": "text",
                    "computed": {
                        "kind": "function",
                        "function": "number_to_words",
                        "source": "total",
                    },
                },
            ]
        },
    )
    assert response.status_code == 422


# ── help_text tests ──────────────────────────────────────────────────────────


@pytest.mark.asyncio
async def test_owner_can_set_help_text(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    """Owner can PATCH help_text and receive the updated value in the response."""
    tpl_id, ver_id = _seed_template_with_variables(
        ["fecha_inicio"], CONFTEST_USER_ID, fake_template_repo, "HelpTextSet"
    )

    response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}/versions/{ver_id}/variables-meta",
        headers=auth_headers,
        json={
            "overrides": [
                {
                    "name": "fecha_inicio",
                    "type": "text",
                    "help_text": "Usar formato DD/MM/YYYY",
                }
            ]
        },
    )
    assert response.status_code == 200
    data = response.json()
    meta_by_name = {m["name"]: m for m in data["variables_meta"]}
    assert meta_by_name["fecha_inicio"]["help_text"] == "Usar formato DD/MM/YYYY"


@pytest.mark.asyncio
async def test_help_text_persists_across_round_trip(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    """Write help_text, then GET the template — the value is preserved."""
    tpl_id, ver_id = _seed_template_with_variables(
        ["campo"], CONFTEST_USER_ID, fake_template_repo, "HelpTextRoundTrip"
    )

    # Write
    await async_client.patch(
        f"/api/v1/templates/{tpl_id}/versions/{ver_id}/variables-meta",
        headers=auth_headers,
        json={
            "overrides": [
                {"name": "campo", "type": "text", "help_text": "Texto de ayuda"}
            ]
        },
    )

    # Read back via GET
    get_response = await async_client.get(
        f"/api/v1/templates/{tpl_id}",
        headers=auth_headers,
    )
    assert get_response.status_code == 200
    get_data = get_response.json()
    version_data = next(v for v in get_data["versions"] if v["id"] == str(ver_id))
    meta_by_name = {m["name"]: m for m in version_data["variables_meta"]}
    assert meta_by_name["campo"]["help_text"] == "Texto de ayuda"


@pytest.mark.asyncio
async def test_help_text_can_be_cleared_with_null(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    """Setting help_text to null clears a previously set value."""
    tpl_id, ver_id = _seed_template_with_variables(
        ["campo"], CONFTEST_USER_ID, fake_template_repo, "HelpTextClear"
    )

    # First set a non-null value
    await async_client.patch(
        f"/api/v1/templates/{tpl_id}/versions/{ver_id}/variables-meta",
        headers=auth_headers,
        json={
            "overrides": [
                {"name": "campo", "type": "text", "help_text": "Valor inicial"}
            ]
        },
    )

    # Now clear it with null
    response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}/versions/{ver_id}/variables-meta",
        headers=auth_headers,
        json={
            "overrides": [
                {"name": "campo", "type": "text", "help_text": None}
            ]
        },
    )
    assert response.status_code == 200
    data = response.json()
    meta_by_name = {m["name"]: m for m in data["variables_meta"]}
    assert meta_by_name["campo"]["help_text"] is None


@pytest.mark.asyncio
async def test_existing_meta_without_help_text_field_returns_null(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    """Backward compat: variables seeded without help_text serialize help_text=null."""
    tpl_id, ver_id = _seed_template_with_variables(
        ["legacy"], CONFTEST_USER_ID, fake_template_repo, "HelpTextLegacy"
    )

    # Call PATCH with empty overrides to trigger serialization
    response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}/versions/{ver_id}/variables-meta",
        headers=auth_headers,
        json={"overrides": []},
    )
    assert response.status_code == 200
    data = response.json()
    meta_by_name = {m["name"]: m for m in data["variables_meta"]}
    assert meta_by_name["legacy"]["help_text"] is None


# ── GET /templates/{tid}/versions/{vid}/structure ────────────────────────────


@pytest.mark.asyncio
async def test_get_version_structure_without_auth_returns_401(async_client, app):
    from app.presentation.middleware.tenant import get_current_user

    original = app.dependency_overrides.pop(get_current_user, None)
    try:
        fake_id = uuid.uuid4()
        response = await async_client.get(
            f"/api/v1/templates/{fake_id}/versions/{fake_id}/structure"
        )
        assert response.status_code == 401
    finally:
        if original is not None:
            app.dependency_overrides[get_current_user] = original


@pytest.mark.asyncio
async def test_get_version_structure_returns_200_for_owner(
    async_client,
    auth_headers,
    fake_template_repo: FakeTemplateRepository,
    fake_template_engine,
    fake_storage,
):
    """Owner gets the structure dict back as TemplateStructureResponse."""
    # Seed a template + version owned by the test user
    template_id, version_id = _seed_template_with_variables(
        ["nombre"], CONFTEST_USER_ID, fake_template_repo, "StructureTpl"
    )
    # The endpoint downloads bytes from MinIO before passing to the engine —
    # the FakeStorageService raises if the path was never seeded.
    version = fake_template_repo._versions[version_id]
    fake_storage.files[("templates", version.minio_path)] = b"fake-docx-bytes"

    # Configure the engine response
    fake_template_engine.structure_to_return = {
        "headers": [],
        "body": [
            {
                "kind": "paragraph",
                "level": 0,
                "spans": [
                    {"text": "Estimado ", "variable": None},
                    {"text": "{{ nombre }}", "variable": "nombre"},
                    {"text": ", gracias.", "variable": None},
                ],
            }
        ],
        "footers": [],
    }

    response = await async_client.get(
        f"/api/v1/templates/{template_id}/versions/{version_id}/structure",
        headers=auth_headers,
    )

    assert response.status_code == 200
    data = response.json()
    assert data["headers"] == []
    assert data["footers"] == []
    assert len(data["body"]) == 1
    assert data["body"][0]["kind"] == "paragraph"
    assert data["body"][0]["spans"][1] == {
        "text": "{{ nombre }}",
        "variable": "nombre",
    }

    # Reset shared fixture so subsequent tests aren't surprised
    fake_template_engine.structure_to_return = {"headers": [], "body": [], "footers": []}


@pytest.mark.asyncio
async def test_get_version_structure_returns_404_when_version_does_not_exist(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    template_id, _ = _seed_template_with_variables(
        ["x"], CONFTEST_USER_ID, fake_template_repo, "StructureTpl404"
    )
    nonexistent_version = uuid.uuid4()

    response = await async_client.get(
        f"/api/v1/templates/{template_id}/versions/{nonexistent_version}/structure",
        headers=auth_headers,
    )

    assert response.status_code == 404


@pytest.mark.asyncio
async def test_update_template_without_auth_returns_401(async_client, app):
    from app.presentation.middleware.tenant import get_current_user

    original = app.dependency_overrides.pop(get_current_user, None)
    try:
        fake_id = uuid.uuid4()
        response = await async_client.patch(
            f"/api/v1/templates/{fake_id}", json={"name": "New name"}
        )
        assert response.status_code == 401
    finally:
        if original is not None:
            app.dependency_overrides[get_current_user] = original


@pytest.mark.asyncio
async def test_owner_can_rename_template_returns_updated_name_and_owner_name(
    async_client,
    app,
    auth_headers,
    fake_template_repo: FakeTemplateRepository,
    fake_user_repo: FakeUserRepository,
):
    """A non-admin owner can PATCH the name; response includes owner_name."""
    owner_id = uuid.UUID("99999999-9999-9999-9999-999999999999")
    owner = await _ensure_user(
        fake_user_repo, owner_id, "patch-owner@example.com", role="user"
    )
    tpl_id = _upload_template_as_user("PatchOwnerOriginal", owner.id, fake_template_repo)

    owner_current_user = CurrentUser(
        user_id=owner.id,
        tenant_id=CONFTEST_TENANT_ID,
        role="user",
    )

    async def override_as_owner():
        return owner_current_user

    original = app.dependency_overrides.get(get_current_user)
    app.dependency_overrides[get_current_user] = override_as_owner
    try:
        response = await async_client.patch(
            f"/api/v1/templates/{tpl_id}",
            headers=auth_headers,
            json={"name": "PatchOwnerRenamed"},
        )
        assert response.status_code == 200
        data = response.json()
        assert data["name"] == "PatchOwnerRenamed"
        assert data["owner_name"] == owner.full_name
    finally:
        if original is not None:
            app.dependency_overrides[get_current_user] = original
        else:
            app.dependency_overrides.pop(get_current_user, None)


@pytest.mark.asyncio
async def test_update_template_foreign_non_admin_gets_403(
    async_client, app, auth_headers, fake_template_repo: FakeTemplateRepository
):
    """A user who neither owns nor administers the template gets 403."""
    tpl_id = _upload_template_as_user(
        "PatchForeignDenied", CONFTEST_USER_ID, fake_template_repo
    )

    user_b = CurrentUser(
        user_id=USER_B_ID,
        tenant_id=CONFTEST_TENANT_ID,
        role="document_generator",
    )

    async def override_as_user_b():
        return user_b

    original = app.dependency_overrides.get(get_current_user)
    app.dependency_overrides[get_current_user] = override_as_user_b
    try:
        response = await async_client.patch(
            f"/api/v1/templates/{tpl_id}",
            headers=auth_headers,
            json={"name": "ShouldNotApply"},
        )
        assert response.status_code == 403
    finally:
        if original is not None:
            app.dependency_overrides[get_current_user] = original
        else:
            app.dependency_overrides.pop(get_current_user, None)


@pytest.mark.asyncio
async def test_update_template_unknown_id_returns_404(async_client, auth_headers):
    fake_id = uuid.uuid4()
    response = await async_client.patch(
        f"/api/v1/templates/{fake_id}",
        headers=auth_headers,
        json={"name": "Whatever"},
    )
    assert response.status_code == 404


@pytest.mark.asyncio
async def test_update_template_collision_returns_409(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    """Renaming to a name already taken within the tenant returns 409."""
    _upload_template_as_user("PatchTakenName", CONFTEST_USER_ID, fake_template_repo)
    tpl_id = _upload_template_as_user(
        "PatchOriginalName", CONFTEST_USER_ID, fake_template_repo
    )

    response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}",
        headers=auth_headers,
        json={"name": "PatchTakenName"},
    )
    assert response.status_code == 409


@pytest.mark.asyncio
async def test_update_template_empty_body_returns_422(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    tpl_id = _upload_template_as_user(
        "PatchEmptyBody", CONFTEST_USER_ID, fake_template_repo
    )

    response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}", headers=auth_headers, json={}
    )
    assert response.status_code == 422


@pytest.mark.asyncio
async def test_update_template_extra_fields_returns_422(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    tpl_id = _upload_template_as_user(
        "PatchExtraFields", CONFTEST_USER_ID, fake_template_repo
    )

    response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}",
        headers=auth_headers,
        json={"name": "Fine", "not_a_real_field": "oops"},
    )
    assert response.status_code == 422


@pytest.mark.asyncio
async def test_update_template_empty_after_strip_name_returns_422(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    tpl_id = _upload_template_as_user(
        "PatchWhitespaceName", CONFTEST_USER_ID, fake_template_repo
    )

    response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}",
        headers=auth_headers,
        json={"name": "   "},
    )
    assert response.status_code == 422


@pytest.mark.asyncio
async def test_update_template_description_can_be_cleared_with_explicit_null(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    """Setting description to null clears a previously set value — mirrors
    the variables-meta help_text clearing semantics."""
    tpl_id = _upload_template_as_user(
        "PatchDescriptionClear", CONFTEST_USER_ID, fake_template_repo
    )
    fake_template_repo._templates[tpl_id].description = "Valor inicial"

    response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}",
        headers=auth_headers,
        json={"description": None},
    )
    assert response.status_code == 200
    data = response.json()
    assert data["description"] is None
    assert fake_template_repo._templates[tpl_id].description is None


@pytest.mark.asyncio
async def test_update_template_description_only_null_body_is_valid(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    """A body of just `{"description": null}` is a valid, non-empty PATCH —
    it explicitly clears the description rather than 422ing as an empty body."""
    tpl_id = _upload_template_as_user(
        "PatchDescriptionOnlyNull", CONFTEST_USER_ID, fake_template_repo
    )

    response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}",
        headers=auth_headers,
        json={"description": None},
    )
    assert response.status_code == 200


@pytest.mark.asyncio
async def test_update_template_name_explicit_null_returns_422(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    """`{"name": null}` must 422 — unlike description, name can never be
    cleared (it has a min-length requirement) and must not silently no-op."""
    tpl_id = _upload_template_as_user(
        "PatchNameExplicitNull", CONFTEST_USER_ID, fake_template_repo
    )

    response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}",
        headers=auth_headers,
        json={"name": None},
    )
    assert response.status_code == 422


# ── owner_name / total in list response ──────────────────────────────────────


@pytest.mark.asyncio
async def test_list_response_includes_owner_name_and_total(
    async_client,
    auth_headers,
    fake_template_repo: FakeTemplateRepository,
    fake_user_repo: FakeUserRepository,
):
    """Each list row carries owner_name; total reflects the accessible count."""
    owner_id = uuid.UUID("66666666-6666-6666-6666-666666666666")
    owner = await _ensure_user(
        fake_user_repo, owner_id, "list-owner-name@example.com", role="user"
    )
    _upload_template_as_user("ListOwnerNameTpl", owner.id, fake_template_repo)

    response = await async_client.get(
        "/api/v1/templates",
        headers=auth_headers,
        params={"search": "ListOwnerNameTpl"},
    )
    assert response.status_code == 200
    data = response.json()
    assert data["total"] == len(data["items"])
    assert len(data["items"]) >= 1
    for item in data["items"]:
        assert "owner_name" in item
    match = next(i for i in data["items"] if i["name"] == "ListOwnerNameTpl")
    assert match["owner_name"] == owner.full_name


@pytest.mark.asyncio
async def test_get_version_structure_returns_404_on_path_mismatch(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    """Version belongs to template A but URL says template B → 404 (prevents oracle leakage)."""
    _, version_a = _seed_template_with_variables(
        ["a"], CONFTEST_USER_ID, fake_template_repo, "StructureTplA"
    )
    template_b, _ = _seed_template_with_variables(
        ["b"], CONFTEST_USER_ID, fake_template_repo, "StructureTplB"
    )

    response = await async_client.get(
        f"/api/v1/templates/{template_b}/versions/{version_a}/structure",
        headers=auth_headers,
    )

    assert response.status_code == 404


# ── folder_id integration (GET /templates filter + PATCH assignment) ─────────

FOLDER_OWNER_ID = uuid.UUID("55555555-5555-5555-5555-555555555555")
FOLDER_STRANGER_ID = uuid.UUID("44444444-4444-4444-4444-444444444444")


async def _create_folder_via_api(async_client, auth_headers, name: str) -> str:
    response = await async_client.post(
        "/api/v1/folders", headers=auth_headers, json={"name": name}
    )
    assert response.status_code == 201, response.text
    return response.json()["id"]


@pytest.mark.asyncio
async def test_template_response_includes_folder_id_field(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    tpl_id = _upload_template_as_user(
        "FolderIdFieldPresent", CONFTEST_USER_ID, fake_template_repo
    )
    response = await async_client.get(f"/api/v1/templates/{tpl_id}", headers=auth_headers)
    assert response.status_code == 200
    assert "folder_id" in response.json()
    assert response.json()["folder_id"] is None


@pytest.mark.asyncio
async def test_owner_can_assign_template_to_own_folder(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    folder_id = await _create_folder_via_api(async_client, auth_headers, "OwnFolderAssign")
    tpl_id = _upload_template_as_user(
        "AssignToOwnFolder", CONFTEST_USER_ID, fake_template_repo
    )

    response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}",
        headers=auth_headers,
        json={"folder_id": folder_id},
    )
    assert response.status_code == 200
    assert response.json()["folder_id"] == folder_id


@pytest.mark.asyncio
async def test_assign_to_foreign_folder_returns_404_non_leaking(
    async_client, app, auth_headers, fake_template_repo: FakeTemplateRepository
):
    """A folder owned by another user cannot be used as a target — 404, and
    the error must not reveal anything about the foreign folder."""
    stranger = CurrentUser(
        user_id=FOLDER_STRANGER_ID,
        tenant_id=CONFTEST_TENANT_ID,
        role="document_generator",
    )

    async def override_as_stranger():
        return stranger

    original = app.dependency_overrides.get(get_current_user)
    app.dependency_overrides[get_current_user] = override_as_stranger
    try:
        foreign_folder_id = await _create_folder_via_api(
            async_client, auth_headers, "StrangerOwnedFolder"
        )
    finally:
        if original is not None:
            app.dependency_overrides[get_current_user] = original
        else:
            app.dependency_overrides.pop(get_current_user, None)

    tpl_id = _upload_template_as_user(
        "AssignToForeignFolder", CONFTEST_USER_ID, fake_template_repo
    )

    response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}",
        headers=auth_headers,
        json={"folder_id": foreign_folder_id},
    )
    assert response.status_code == 404
    assert "Stranger" not in response.text


@pytest.mark.asyncio
async def test_non_owner_admin_cannot_refile_someone_elses_template(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    """Admins can rename another user's template but must NOT be able to
    re-file it into a folder — folders are strictly personal."""
    folder_id = await _create_folder_via_api(
        async_client, auth_headers, "AdminOwnFolderForRefileTest"
    )
    tpl_id = _upload_template_as_user(
        "AdminCannotRefile", USER_B_ID, fake_template_repo
    )

    # Conftest user is admin but not the owner of this template.
    response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}",
        headers=auth_headers,
        json={"folder_id": folder_id},
    )
    assert response.status_code == 403


@pytest.mark.asyncio
async def test_owner_can_unfile_template_with_explicit_null(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    folder_id = await _create_folder_via_api(async_client, auth_headers, "UnfileMeFolder")
    tpl_id = _upload_template_as_user("ToBeUnfiled", CONFTEST_USER_ID, fake_template_repo)

    assign_response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}",
        headers=auth_headers,
        json={"folder_id": folder_id},
    )
    assert assign_response.json()["folder_id"] == folder_id

    unfile_response = await async_client.patch(
        f"/api/v1/templates/{tpl_id}",
        headers=auth_headers,
        json={"folder_id": None},
    )
    assert unfile_response.status_code == 200
    assert unfile_response.json()["folder_id"] is None


@pytest.mark.asyncio
async def test_list_templates_filter_by_folder_id_returns_only_that_folder(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    folder_a = await _create_folder_via_api(async_client, auth_headers, "FilterFolderA")
    folder_b = await _create_folder_via_api(async_client, auth_headers, "FilterFolderB")

    tpl_in_a = _upload_template_as_user(
        "InFilterFolderA", CONFTEST_USER_ID, fake_template_repo
    )
    tpl_in_b = _upload_template_as_user(
        "InFilterFolderB", CONFTEST_USER_ID, fake_template_repo
    )

    await async_client.patch(
        f"/api/v1/templates/{tpl_in_a}", headers=auth_headers, json={"folder_id": folder_a}
    )
    await async_client.patch(
        f"/api/v1/templates/{tpl_in_b}", headers=auth_headers, json={"folder_id": folder_b}
    )

    response = await async_client.get(
        "/api/v1/templates", headers=auth_headers, params={"folder_id": folder_a}
    )
    assert response.status_code == 200
    ids = [item["id"] for item in response.json()["items"]]
    assert str(tpl_in_a) in ids
    assert str(tpl_in_b) not in ids


@pytest.mark.asyncio
async def test_list_templates_folder_id_none_returns_only_unfiled(
    async_client, auth_headers, fake_template_repo: FakeTemplateRepository
):
    folder_id = await _create_folder_via_api(async_client, auth_headers, "UnfiledFilterFolder")

    # Unique shared search term: the fake template repo is session-scoped and
    # shared across the whole test file, so without narrowing via `search`
    # the default page size (20) could push these templates past page 1 in a
    # full suite run — same caveat documented on
    # test_upload_template_appears_in_list above.
    unique_prefix = f"UnfiledOnlyProbe{uuid.uuid4().hex[:8]}"
    tpl_filed = _upload_template_as_user(
        f"{unique_prefix}Filed", CONFTEST_USER_ID, fake_template_repo
    )
    tpl_unfiled = _upload_template_as_user(
        f"{unique_prefix}Unfiled", CONFTEST_USER_ID, fake_template_repo
    )
    await async_client.patch(
        f"/api/v1/templates/{tpl_filed}", headers=auth_headers, json={"folder_id": folder_id}
    )

    response = await async_client.get(
        "/api/v1/templates",
        headers=auth_headers,
        params={"folder_id": "none", "search": unique_prefix},
    )
    assert response.status_code == 200
    ids = [item["id"] for item in response.json()["items"]]
    assert str(tpl_unfiled) in ids
    assert str(tpl_filed) not in ids


@pytest.mark.asyncio
async def test_folder_filter_does_not_leak_shared_templates_across_owners(
    async_client, app, auth_headers, fake_template_repo: FakeTemplateRepository
):
    """Regression: user B shares a template with user A (the conftest admin
    user). User A creates their own folder and filters by it — user B's
    shared template (even if user B also has a folder with the same
    conceptual grouping) must NEVER surface just because it's accessible to
    user A. The folder filter must intersect with owned+shared visibility,
    never bypass it."""
    # User A (conftest admin) creates their own folder.
    folder_a = await _create_folder_via_api(
        async_client, auth_headers, "LeakRegressionFolderA"
    )

    # User B owns and shares a template with user A.
    shared_tpl_id = _upload_template_as_user(
        "LeakRegressionSharedTpl", USER_B_ID, fake_template_repo
    )
    await fake_template_repo.add_share(
        template_id=shared_tpl_id,
        user_id=CONFTEST_USER_ID,
        tenant_id=CONFTEST_TENANT_ID,
        shared_by=USER_B_ID,
    )

    # A template genuinely filed in user A's own folder.
    own_filed_tpl_id = _upload_template_as_user(
        "LeakRegressionOwnFiled", CONFTEST_USER_ID, fake_template_repo
    )
    await async_client.patch(
        f"/api/v1/templates/{own_filed_tpl_id}",
        headers=auth_headers,
        json={"folder_id": folder_a},
    )

    response = await async_client.get(
        "/api/v1/templates", headers=auth_headers, params={"folder_id": folder_a}
    )
    assert response.status_code == 200
    ids = [item["id"] for item in response.json()["items"]]
    assert str(own_filed_tpl_id) in ids
    assert str(shared_tpl_id) not in ids


@pytest.mark.asyncio
async def test_folder_id_none_filter_does_not_leak_unshared_unfiled_templates(
    async_client, app, auth_headers, fake_template_repo: FakeTemplateRepository
):
    """Regression for the `folder_id=none` (unfiled) branch specifically.

    User B owns an UNFILED template in the same tenant that is NOT shared
    with user A. User A (a non-admin `document_generator`, NOT the
    conftest admin — admins legitimately see every tenant template via
    `can_view_all_templates`, so this scenario must use a role that is
    actually subject to owned/shared visibility) lists with
    `folder_id=none` (their own unfiled templates) — user B's unfiled
    template must never appear.

    Reasoning on the code path (why this test would catch a regression):
    in `SQLAlchemyTemplateRepository.list_accessible`
    (template_repository.py), the owned/shared accessible-ids filter is
    applied to `stmt`/`count_stmt` first (building the visibility scope),
    and the `folder_filter_unfiled` clause (`TemplateModel.folder_id.is_(None)`)
    is then `.where(...)`-chained onto that SAME statement — so the two
    conditions are ANDed together. If the unfiled branch were instead built
    as a standalone query (e.g. `select(TemplateModel).where(folder_id.is_(None))`
    without also filtering by owned/shared ids), it would return every
    unfiled template in the tenant regardless of ownership, and user B's
    unfiled template would leak into user A's result — which is exactly
    what this test asserts never happens. The fake repository mirrors this
    by filtering `items` down to owned+shared FIRST, then applying the
    unfiled filter on that already-narrowed list (see
    `FakeTemplateRepository.list_accessible`).
    """
    unique_prefix = f"UnfiledLeakProbe{uuid.uuid4().hex[:8]}"

    user_a_id = uuid.uuid4()

    # User B's unfiled template — never shared with user A.
    other_owner_unfiled_tpl_id = _upload_template_as_user(
        f"{unique_prefix}OtherOwnerUnfiled", USER_B_ID, fake_template_repo
    )

    # User A's own unfiled template, for a positive control.
    own_unfiled_tpl_id = _upload_template_as_user(
        f"{unique_prefix}OwnUnfiled", user_a_id, fake_template_repo
    )

    user_a = CurrentUser(
        user_id=user_a_id,
        tenant_id=CONFTEST_TENANT_ID,
        role="document_generator",
    )

    async def override_as_user_a():
        return user_a

    original = app.dependency_overrides.get(get_current_user)
    app.dependency_overrides[get_current_user] = override_as_user_a
    try:
        response = await async_client.get(
            "/api/v1/templates",
            headers=auth_headers,
            params={"folder_id": "none", "search": unique_prefix},
        )
    finally:
        if original is not None:
            app.dependency_overrides[get_current_user] = original
        else:
            app.dependency_overrides.pop(get_current_user, None)

    assert response.status_code == 200
    ids = [item["id"] for item in response.json()["items"]]
    assert str(own_unfiled_tpl_id) in ids
    assert str(other_owner_unfiled_tpl_id) not in ids
