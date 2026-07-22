"""
Unit tests for DocxTemplateEngine.render() — sandbox containment + fidelity.

Security invariant (SSTI -> RCE):
    docxtpl builds a PLAIN (non-sandboxed) jinja2.Environment by default, so an
    uploaded .docx whose text contains a Jinja expression like
        {{ cycler.__init__.__globals__.os.popen('id').read() }}
    is EVALUATED server-side at render/preview/generate time -> arbitrary
    command execution. The engine must render every template through a
    jinja2.sandbox.SandboxedEnvironment so that attribute access into Python
    internals raises jinja2.exceptions.SecurityError, which the engine maps to
    the non-leaking domain error TemplateRenderError (a 4xx at the API layer),
    NOT an uncaught 500.

Fidelity invariant:
    Legitimate templates (plain variables, filters, conditionals, missing
    variables rendered as blanks) must keep rendering EXACTLY as before.
"""
import io

import pytest
from docx import Document
from jinja2.exceptions import SecurityError

from app.domain.exceptions import TemplateRenderError
from app.infrastructure.templating.docxtpl_engine import DocxTemplateEngine


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_docx_bytes(paragraphs: list[str]) -> bytes:
    """Build a minimal .docx in memory with one run per paragraph."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph().add_run(text)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _body_text(rendered: bytes) -> str:
    """Extract the concatenated body paragraph text of a rendered .docx."""
    doc = Document(io.BytesIO(rendered))
    return "\n".join(p.text for p in doc.paragraphs)


# The canonical SSTI -> RCE payload. A dummy `{{ x }}` variable is included so
# the upload validator's "no variables" guard would not fire on this template.
_EXPLOIT_PAYLOAD = "{{ cycler.__init__.__globals__.os.popen('id').read() }}"


# ---------------------------------------------------------------------------
# Security: the sandbox must block SSTI -> RCE
# ---------------------------------------------------------------------------


class TestRenderSandboxContainment:
    async def test_ssti_payload_is_rejected_and_no_command_runs(self):
        """The exploit expression must NOT be evaluated.

        With the plain (vulnerable) environment the template renders
        successfully and the shell command output leaks into the document.
        With the SandboxedEnvironment the attribute access into __init__ /
        __globals__ raises SecurityError, which the engine maps to
        TemplateRenderError. Asserting that TemplateRenderError is raised
        proves the command never produced a rendered result.
        """
        engine = DocxTemplateEngine()
        raw = _make_docx_bytes([_EXPLOIT_PAYLOAD, "Hello {{ x }}"])

        with pytest.raises(TemplateRenderError):
            await engine.render(raw, {"x": "1"})

    async def test_ssti_rejection_is_caused_by_a_sandbox_security_error(self):
        """The mapped domain error must chain the underlying SecurityError.

        This documents the containment mechanism: the expression was stopped
        by the Jinja sandbox at evaluation time, not merely rejected by a
        static string check.
        """
        engine = DocxTemplateEngine()
        raw = _make_docx_bytes([_EXPLOIT_PAYLOAD, "Hello {{ x }}"])

        with pytest.raises(TemplateRenderError) as exc_info:
            await engine.render(raw, {"x": "1"})

        assert isinstance(exc_info.value.__cause__, SecurityError)

    async def test_class_mro_probe_is_rejected(self):
        """Another common SSTI probe (__class__ walking) is also contained."""
        engine = DocxTemplateEngine()
        raw = _make_docx_bytes(["{{ x.__class__.__mro__ }}", "Hi {{ x }}"])

        with pytest.raises(TemplateRenderError):
            await engine.render(raw, {"x": "1"})


# ---------------------------------------------------------------------------
# Fidelity: legitimate templates keep rendering identically
# ---------------------------------------------------------------------------


class TestRenderLegitimateTemplates:
    async def test_plain_variables_render_unchanged(self):
        engine = DocxTemplateEngine()
        raw = _make_docx_bytes(["Hola {{ nombre }}, el monto es {{ monto }}."])

        rendered = await engine.render(raw, {"nombre": "Ada", "monto": "100"})

        assert _body_text(rendered) == "Hola Ada, el monto es 100."

    async def test_jinja_filter_and_conditional_still_work(self):
        """SandboxedEnvironment must NOT break normal Jinja constructs."""
        engine = DocxTemplateEngine()
        raw = _make_docx_bytes(
            ["{{ nombre | upper }}{% if activo %} (activo){% endif %}"]
        )

        rendered = await engine.render(raw, {"nombre": "ada", "activo": True})

        assert _body_text(rendered) == "ADA (activo)"

    async def test_missing_variable_renders_as_blank(self):
        """docxtpl default Undefined behavior (blank) must be preserved."""
        engine = DocxTemplateEngine()
        raw = _make_docx_bytes(["Titular: {{ falta }}."])

        rendered = await engine.render(raw, {})

        assert _body_text(rendered) == "Titular: ."

    async def test_xml_special_chars_are_escaped(self):
        """autoescape must still protect against XML injection from values."""
        engine = DocxTemplateEngine()
        raw = _make_docx_bytes(["Valor: {{ v }}"])

        rendered = await engine.render(raw, {"v": "<b>&</b>"})

        # The value must appear as literal text, not as injected XML markup.
        assert _body_text(rendered) == "Valor: <b>&</b>"
