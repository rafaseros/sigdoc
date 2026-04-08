"""
Unit tests for DocxTemplateEngine.auto_fix() — specifically verifies the
paragraph-level fix that handles variables split across multiple Word runs.

Word applies formatting (bold, italic, font changes) at run boundaries, which
means a variable like {{ UserName }} can be stored as three separate runs:
  Run 1: "{{ "
  Run 2: "UserName"   (formatted differently)
  Run 3: " }}"
The old run-level regex never saw a complete pattern, so the file came back
unchanged. The new implementation works with para.text (the full concatenated
text) and then rebuilds the runs.
"""
import io

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from app.infrastructure.templating.docxtpl_engine import DocxTemplateEngine


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_docx_bytes(paragraphs: list[list[str]]) -> bytes:
    """
    Build a minimal .docx in memory.

    ``paragraphs`` is a list of paragraphs; each paragraph is a list of run
    text fragments that will be added as *separate runs* to simulate what Word
    does when formatting spans part of a variable marker.
    """
    doc = Document()
    for run_texts in paragraphs:
        para = doc.add_paragraph()
        for text in run_texts:
            para.add_run(text)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _make_docx_with_table(cell_run_texts: list[str]) -> bytes:
    """Build a .docx with a 1x1 table whose only cell has split runs."""
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    para = cell.paragraphs[0]
    for text in cell_run_texts:
        para.add_run(text)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _read_all_paragraph_texts(docx_bytes: bytes) -> list[str]:
    """Return the full text of every paragraph (body + tables) in a docx."""
    doc = Document(io.BytesIO(docx_bytes))
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        texts.append(para.text)
    return texts


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


class TestAutoFixSplitRuns:
    """Variables spread across multiple runs must still be fixed."""

    async def test_fixes_uppercase_split_across_three_runs(self):
        """
        '{{ ', 'UserName', ' }}' in three separate runs must become
        '{{ user_name }}' in the first run.
        """
        engine = DocxTemplateEngine()
        raw = _make_docx_bytes([["{{ ", "UserName", " }}"]])

        fixed = await engine.auto_fix(raw)

        texts = _read_all_paragraph_texts(fixed)
        assert any("{{ user_name }}" in t for t in texts), (
            f"Expected '{{{{ user_name }}}}' in output paragraphs, got: {texts}"
        )

    async def test_fixes_camel_case_split_across_runs(self):
        """camelCase variable split across runs → snake_case after fix."""
        engine = DocxTemplateEngine()
        raw = _make_docx_bytes([["{{", "firstName", "}}"]])

        fixed = await engine.auto_fix(raw)

        texts = _read_all_paragraph_texts(fixed)
        assert any("{{ first_name }}" in t for t in texts), (
            f"Expected '{{{{ first_name }}}}' in output, got: {texts}"
        )

    async def test_fixes_no_space_variable_split_across_runs(self):
        """
        '{{', 'username', '}}' (no spaces, lowercase) → '{{ username }}'.
        """
        engine = DocxTemplateEngine()
        raw = _make_docx_bytes([["{{", "username", "}}"]])

        fixed = await engine.auto_fix(raw)

        texts = _read_all_paragraph_texts(fixed)
        assert any("{{ username }}" in t for t in texts), (
            f"Expected '{{{{ username }}}}' in output, got: {texts}"
        )

    async def test_fixes_pascal_case_in_single_run(self):
        """Control case: single-run PascalCase variable is also fixed."""
        engine = DocxTemplateEngine()
        raw = _make_docx_bytes([["{{ CompanyName }}"]])

        fixed = await engine.auto_fix(raw)

        texts = _read_all_paragraph_texts(fixed)
        assert any("{{ company_name }}" in t for t in texts), (
            f"Expected '{{{{ company_name }}}}' in output, got: {texts}"
        )

    async def test_correct_variable_unchanged(self):
        """A correctly formatted variable must survive auto_fix unchanged."""
        engine = DocxTemplateEngine()
        raw = _make_docx_bytes([["{{ first_name }}"]])

        fixed = await engine.auto_fix(raw)

        texts = _read_all_paragraph_texts(fixed)
        assert any("{{ first_name }}" in t for t in texts), (
            f"Expected '{{{{ first_name }}}}' to be preserved, got: {texts}"
        )

    async def test_fixes_variable_in_table_cell_split_runs(self):
        """Variables split across runs inside table cells must also be fixed."""
        engine = DocxTemplateEngine()
        raw = _make_docx_with_table(["{{ ", "InvoiceNumber", " }}"])

        fixed = await engine.auto_fix(raw)

        texts = _read_all_paragraph_texts(fixed)
        assert any("{{ invoice_number }}" in t for t in texts), (
            f"Expected '{{{{ invoice_number }}}}' in table cell, got: {texts}"
        )

    async def test_multiple_paragraphs_each_with_split_runs(self):
        """Multiple paragraphs with split-run variables are all fixed."""
        engine = DocxTemplateEngine()
        raw = _make_docx_bytes([
            ["{{ ", "FirstName", " }}"],
            ["{{ ", "LastName", " }}"],
            ["{{ ", "CompanyName", " }}"],
        ])

        fixed = await engine.auto_fix(raw)

        texts = _read_all_paragraph_texts(fixed)
        joined = " ".join(texts)
        assert "{{ first_name }}" in joined, f"first_name missing in: {texts}"
        assert "{{ last_name }}" in joined, f"last_name missing in: {texts}"
        assert "{{ company_name }}" in joined, f"company_name missing in: {texts}"

    async def test_preserves_surrounding_text(self):
        """
        Plain text surrounding the variable must not be mangled.
        E.g. 'Dear {{ ', 'UserName', ' }}, welcome!' → 'Dear {{ user_name }}, welcome!'
        """
        engine = DocxTemplateEngine()
        raw = _make_docx_bytes([["Dear {{ ", "UserName", " }}, welcome!"]])

        fixed = await engine.auto_fix(raw)

        texts = _read_all_paragraph_texts(fixed)
        joined = " ".join(texts)
        assert "{{ user_name }}" in joined, f"Variable not fixed in: {texts}"
        assert "Dear" in joined, f"'Dear' prefix lost in: {texts}"
        assert "welcome!" in joined, f"'welcome!' suffix lost in: {texts}"

    async def test_returns_bytes(self):
        """auto_fix must always return bytes (not None or empty)."""
        engine = DocxTemplateEngine()
        raw = _make_docx_bytes([["{{ user_name }}"]])

        fixed = await engine.auto_fix(raw)

        assert isinstance(fixed, bytes)
        assert len(fixed) > 0
