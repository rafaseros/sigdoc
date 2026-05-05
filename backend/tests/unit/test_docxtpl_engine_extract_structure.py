"""
Unit tests for DocxTemplateEngine.extract_structure() — full document structure
preview used by the generation UI.

Each paragraph is returned as a node with `kind`, `level`, and a list of `spans`
where each span is either plain text or a {{ variable }} placeholder. Headers,
body and footers are returned as three separate lists.
"""

import io

import pytest
from docx import Document

from app.infrastructure.templating.docxtpl_engine import DocxTemplateEngine


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _doc_bytes(build) -> bytes:
    """Run `build(doc)` on a fresh Document and return the bytes."""
    doc = Document()
    build(doc)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


class TestExtractStructureBody:
    """Body section: paragraphs, headings, placeholders, edge cases."""

    async def test_plain_paragraph_returns_single_text_span(self):
        engine = DocxTemplateEngine()
        raw = _doc_bytes(lambda d: d.add_paragraph("Hola mundo"))

        result = await engine.extract_structure(raw)

        assert result["body"] == [
            {"kind": "paragraph", "level": 0, "spans": [{"text": "Hola mundo", "variable": None}]}
        ]
        assert result["headers"] == []
        assert result["footers"] == []

    async def test_paragraph_with_placeholder_splits_into_spans(self):
        engine = DocxTemplateEngine()
        raw = _doc_bytes(
            lambda d: d.add_paragraph("Estimado {{ nombre }}, gracias.")
        )

        result = await engine.extract_structure(raw)

        body = result["body"]
        assert len(body) == 1
        assert body[0]["kind"] == "paragraph"
        assert body[0]["spans"] == [
            {"text": "Estimado ", "variable": None},
            {"text": "{{ nombre }}", "variable": "nombre"},
            {"text": ", gracias.", "variable": None},
        ]

    async def test_paragraph_starting_with_placeholder_no_leading_text_span(self):
        """When the placeholder is at position 0, no empty text span is emitted."""
        engine = DocxTemplateEngine()
        raw = _doc_bytes(lambda d: d.add_paragraph("{{ nombre }} firma."))

        result = await engine.extract_structure(raw)

        spans = result["body"][0]["spans"]
        assert spans[0] == {"text": "{{ nombre }}", "variable": "nombre"}
        assert spans[1] == {"text": " firma.", "variable": None}
        assert len(spans) == 2

    async def test_paragraph_ending_with_placeholder_no_trailing_text_span(self):
        engine = DocxTemplateEngine()
        raw = _doc_bytes(lambda d: d.add_paragraph("Firma: {{ nombre }}"))

        result = await engine.extract_structure(raw)

        spans = result["body"][0]["spans"]
        assert spans == [
            {"text": "Firma: ", "variable": None},
            {"text": "{{ nombre }}", "variable": "nombre"},
        ]

    async def test_multiple_placeholders_same_paragraph(self):
        engine = DocxTemplateEngine()
        raw = _doc_bytes(
            lambda d: d.add_paragraph("Entre {{ empresa }} y {{ contratado }}.")
        )

        result = await engine.extract_structure(raw)

        spans = result["body"][0]["spans"]
        variables = [s["variable"] for s in spans if s["variable"]]
        assert variables == ["empresa", "contratado"]

    async def test_empty_paragraphs_are_skipped(self):
        """Whitespace-only paragraphs do not become nodes — they would just clutter the preview."""
        engine = DocxTemplateEngine()

        def build(d):
            d.add_paragraph("Primero")
            d.add_paragraph("")
            d.add_paragraph("   ")
            d.add_paragraph("Tercero")

        raw = _doc_bytes(build)
        result = await engine.extract_structure(raw)

        texts = [n["spans"][0]["text"] for n in result["body"]]
        assert texts == ["Primero", "Tercero"]

    async def test_heading_detected_with_level(self):
        engine = DocxTemplateEngine()

        def build(d):
            d.add_heading("Título principal", level=1)
            d.add_heading("Sección", level=2)
            d.add_paragraph("Contenido normal")

        raw = _doc_bytes(build)
        result = await engine.extract_structure(raw)

        body = result["body"]
        assert body[0]["kind"] == "heading"
        assert body[0]["level"] == 1
        assert body[1]["kind"] == "heading"
        assert body[1]["level"] == 2
        assert body[2]["kind"] == "paragraph"
        assert body[2]["level"] == 0


class TestExtractStructureHeadersFooters:
    """Headers and footers must be returned in their own lists, not in body."""

    async def test_header_paragraphs_returned_separately(self):
        engine = DocxTemplateEngine()

        def build(d):
            section = d.sections[0]
            section.header.paragraphs[0].text = "Encabezado de página"
            d.add_paragraph("Cuerpo del contrato")

        raw = _doc_bytes(build)
        result = await engine.extract_structure(raw)

        assert len(result["headers"]) == 1
        assert result["headers"][0]["spans"][0]["text"] == "Encabezado de página"
        assert len(result["body"]) == 1
        assert result["body"][0]["spans"][0]["text"] == "Cuerpo del contrato"

    async def test_footer_paragraphs_returned_separately(self):
        engine = DocxTemplateEngine()

        def build(d):
            section = d.sections[0]
            section.footer.paragraphs[0].text = "Página {{ page }}"
            d.add_paragraph("Cuerpo")

        raw = _doc_bytes(build)
        result = await engine.extract_structure(raw)

        assert len(result["footers"]) == 1
        assert result["footers"][0]["spans"] == [
            {"text": "Página ", "variable": None},
            {"text": "{{ page }}", "variable": "page"},
        ]

    async def test_placeholders_in_headers_are_parsed(self):
        engine = DocxTemplateEngine()

        def build(d):
            d.sections[0].header.paragraphs[0].text = (
                "Contrato {{ numero_contrato }} — {{ nombre_empresa }}"
            )
            d.add_paragraph("Cuerpo")

        raw = _doc_bytes(build)
        result = await engine.extract_structure(raw)

        header_vars = [
            s["variable"] for s in result["headers"][0]["spans"] if s["variable"]
        ]
        assert header_vars == ["numero_contrato", "nombre_empresa"]


class TestExtractStructureMisc:
    async def test_returns_three_keys_always(self):
        """Even an empty .docx should return all three keys with empty lists."""
        engine = DocxTemplateEngine()
        raw = _doc_bytes(lambda d: None)

        result = await engine.extract_structure(raw)

        assert result == {"headers": [], "body": [], "footers": []}

    async def test_invalid_variable_name_with_dash_is_not_a_placeholder(self):
        """Variable names allow only [a-zA-Z_][a-zA-Z0-9_]*. A dash makes it plain text."""
        engine = DocxTemplateEngine()
        raw = _doc_bytes(lambda d: d.add_paragraph("Texto {{ no-valid }} fin"))

        result = await engine.extract_structure(raw)

        # Whole paragraph should be a single text span — the regex shouldn't match.
        spans = result["body"][0]["spans"]
        assert len(spans) == 1
        assert spans[0]["variable"] is None


# ─── Lists ───────────────────────────────────────────────────────────────────


class TestExtractStructureLists:
    async def test_bullet_list_detected(self):
        engine = DocxTemplateEngine()

        def build(d):
            d.add_paragraph("Items:", style=None)
            d.add_paragraph("Primero", style="List Bullet")
            d.add_paragraph("Segundo", style="List Bullet")

        raw = _doc_bytes(build)
        result = await engine.extract_structure(raw)

        kinds = [n["kind"] for n in result["body"]]
        # First item is a plain paragraph; the next two are bullet list items.
        assert kinds[0] == "paragraph"
        assert kinds[1] == "list_bullet"
        assert kinds[2] == "list_bullet"
        assert result["body"][1]["level"] == 1
        assert result["body"][1]["spans"][0]["text"] == "Primero"

    async def test_numbered_list_detected(self):
        engine = DocxTemplateEngine()

        def build(d):
            d.add_paragraph("Uno", style="List Number")
            d.add_paragraph("Dos", style="List Number")

        raw = _doc_bytes(build)
        result = await engine.extract_structure(raw)

        assert all(n["kind"] == "list_number" for n in result["body"])
        assert all(n["level"] == 1 for n in result["body"])

    async def test_list_with_placeholder_splits_spans(self):
        engine = DocxTemplateEngine()
        raw = _doc_bytes(
            lambda d: d.add_paragraph("Pago a {{ nombre }}", style="List Bullet")
        )

        result = await engine.extract_structure(raw)

        node = result["body"][0]
        assert node["kind"] == "list_bullet"
        variables = [s["variable"] for s in node["spans"] if s["variable"]]
        assert variables == ["nombre"]


# ─── Tables ──────────────────────────────────────────────────────────────────


class TestExtractStructureTables:
    async def test_simple_table_returned_as_table_node(self):
        engine = DocxTemplateEngine()

        def build(d):
            d.add_paragraph("Antes")
            t = d.add_table(rows=2, cols=2)
            t.rows[0].cells[0].text = "Campo"
            t.rows[0].cells[1].text = "Valor"
            t.rows[1].cells[0].text = "Nombre"
            t.rows[1].cells[1].text = "Acme"
            d.add_paragraph("Despues")

        raw = _doc_bytes(build)
        result = await engine.extract_structure(raw)

        body = result["body"]
        # paragraph - table - paragraph in document order
        assert [n["kind"] for n in body] == ["paragraph", "table", "paragraph"]
        table = body[1]
        assert len(table["rows"]) == 2
        assert len(table["rows"][0]["cells"]) == 2
        # Cell content survives as a nested paragraph node
        first_cell_nodes = table["rows"][0]["cells"][0]["nodes"]
        assert first_cell_nodes[0]["spans"][0]["text"] == "Campo"

    async def test_table_with_placeholder_in_cell(self):
        engine = DocxTemplateEngine()

        def build(d):
            t = d.add_table(rows=1, cols=2)
            t.rows[0].cells[0].text = "Empresa"
            t.rows[0].cells[1].text = "{{ nombre_empresa }}"

        raw = _doc_bytes(build)
        result = await engine.extract_structure(raw)

        cell_nodes = result["body"][0]["rows"][0]["cells"][1]["nodes"]
        spans = cell_nodes[0]["spans"]
        assert any(
            s["variable"] == "nombre_empresa" and s["text"] == "{{ nombre_empresa }}"
            for s in spans
        )

    async def test_blank_table_is_skipped(self):
        """Tables whose every cell is empty are usually layout artifacts —
        we drop them so the preview stays meaningful."""
        engine = DocxTemplateEngine()

        def build(d):
            d.add_paragraph("Antes")
            d.add_table(rows=2, cols=3)  # all cells empty
            d.add_paragraph("Despues")

        raw = _doc_bytes(build)
        result = await engine.extract_structure(raw)

        kinds = [n["kind"] for n in result["body"]]
        assert "table" not in kinds
        assert kinds == ["paragraph", "paragraph"]

    async def test_table_in_header_is_returned_in_headers(self):
        engine = DocxTemplateEngine()

        def build(d):
            section = d.sections[0]
            t = section.header.add_table(rows=1, cols=2, width=10)
            t.rows[0].cells[0].text = "logo"
            t.rows[0].cells[1].text = "Contrato {{ numero_contrato }}"
            d.add_paragraph("cuerpo")

        raw = _doc_bytes(build)
        result = await engine.extract_structure(raw)

        # Header now carries a table node
        assert any(n["kind"] == "table" for n in result["headers"])
        # Body still has the plain paragraph
        assert result["body"][0]["spans"][0]["text"] == "cuerpo"
