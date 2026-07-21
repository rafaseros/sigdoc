"""
Unit tests for DocxTemplateEngine.apply_variable_mappings() — the
template-from-example rewrite.

The engine receives a filled example .docx plus [{"text", "variable"}, ...]
mappings and must replace every literal occurrence with a {{ placeholder }}
WITHOUT altering Word formatting:

- match inside a single run  → split that run's text in place
- match spanning runs        → placeholder inherits the FIRST run's formatting,
                               preserved segments keep their own run formatting
- untouched paragraphs       → runs remain exactly as they were
"""
import io

import pytest
from docx import Document
from docx.shared import Pt

from app.domain.exceptions import (
    InvalidVariableMappingError,
    MappingTextNotFoundError,
)
from app.infrastructure.templating.docxtpl_engine import DocxTemplateEngine


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _doc_to_bytes(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _load(docx_bytes: bytes):
    return Document(io.BytesIO(docx_bytes))


def _body_texts(docx_bytes: bytes) -> list[str]:
    return [p.text for p in _load(docx_bytes).paragraphs]


def _all_texts(docx_bytes: bytes) -> list[str]:
    """Every paragraph text: body, body tables, headers, footers."""
    doc = _load(docx_bytes)
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(p.text for p in cell.paragraphs)
    for section in doc.sections:
        texts.extend(p.text for p in section.header.paragraphs)
        texts.extend(p.text for p in section.footer.paragraphs)
    return texts


# ---------------------------------------------------------------------------
# Formatting preservation
# ---------------------------------------------------------------------------


class TestSingleRunReplacement:
    async def test_split_inside_single_run_preserves_text_and_formatting(self):
        """Match inside one run → that run's text is split in place; the
        placeholder lives in the same run and thus keeps its formatting.
        Surrounding runs are untouched."""
        engine = DocxTemplateEngine()

        doc = Document()
        para = doc.add_paragraph()
        r1 = para.add_run("Entre ")
        r1.italic = True
        r2 = para.add_run("JUAN PÉREZ, con DNI 123")
        r2.bold = True
        r2.font.name = "Arial"
        r2.font.size = Pt(14)
        para.add_run(" firma.")
        raw = _doc_to_bytes(doc)

        result = await engine.apply_variable_mappings(
            raw, [{"text": "JUAN PÉREZ", "variable": "client_name"}]
        )

        out = _load(result)
        p = out.paragraphs[0]
        assert p.text == "Entre {{ client_name }}, con DNI 123 firma."

        # Same number of runs — split happened IN PLACE, no rebuild
        assert len(p.runs) == 3
        assert p.runs[0].text == "Entre "
        assert p.runs[0].italic is True
        assert p.runs[1].text == "{{ client_name }}, con DNI 123"
        assert p.runs[1].bold is True
        assert p.runs[1].font.name == "Arial"
        assert p.runs[1].font.size == Pt(14)
        assert p.runs[2].text == " firma."
        assert p.runs[2].bold is None


class TestSpanningRunsReplacement:
    async def test_match_spanning_runs_inherits_first_run_formatting(self):
        """'JUAN ' (bold) + 'PÉREZ' (plain) → replaced once; the placeholder
        inherits the FIRST run's formatting (bold); rest of paragraph intact."""
        engine = DocxTemplateEngine()

        doc = Document()
        para = doc.add_paragraph()
        para.add_run("Contrato de ")
        r_bold = para.add_run("JUAN ")
        r_bold.bold = True
        para.add_run("PÉREZ")  # plain, different formatting
        para.add_run(" residente en la ciudad.")
        raw = _doc_to_bytes(doc)

        result = await engine.apply_variable_mappings(
            raw, [{"text": "JUAN PÉREZ", "variable": "client_name"}]
        )

        out = _load(result)
        p = out.paragraphs[0]
        assert p.text == "Contrato de {{ client_name }} residente en la ciudad."

        # The placeholder must live in a run with the FIRST matched run's format
        placeholder_runs = [r for r in p.runs if "{{ client_name }}" in r.text]
        assert len(placeholder_runs) == 1
        assert placeholder_runs[0].bold is True

        # Preserved segments keep their own formatting
        assert p.runs[0].text == "Contrato de "
        assert p.runs[0].bold is None
        assert p.runs[-1].text == " residente en la ciudad."
        assert p.runs[-1].bold is None

    async def test_partial_overlap_keeps_prefix_and_suffix_formatting(self):
        """Match starting mid-run and ending mid-run: prefix stays in the first
        run (its formatting), suffix stays in the last run (its formatting)."""
        engine = DocxTemplateEngine()

        doc = Document()
        para = doc.add_paragraph()
        r1 = para.add_run("Sr. JUAN ")
        r1.italic = True
        r2 = para.add_run("PÉREZ firma aquí")
        r2.bold = True
        raw = _doc_to_bytes(doc)

        result = await engine.apply_variable_mappings(
            raw, [{"text": "JUAN PÉREZ", "variable": "client_name"}]
        )

        out = _load(result)
        p = out.paragraphs[0]
        assert p.text == "Sr. {{ client_name }} firma aquí"

        # Prefix 'Sr. ' + placeholder both live in the first (italic) run
        assert p.runs[0].text == "Sr. {{ client_name }}"
        assert p.runs[0].italic is True
        # Suffix keeps the second run's bold formatting
        assert p.runs[1].text == " firma aquí"
        assert p.runs[1].bold is True


# ---------------------------------------------------------------------------
# Coverage: body + tables + headers + footers, all occurrences
# ---------------------------------------------------------------------------


class TestReplacementCoverage:
    async def test_all_occurrences_across_body_table_header_footer(self):
        engine = DocxTemplateEngine()

        doc = Document()
        doc.add_paragraph("Cliente: ACME en el cuerpo. Repetido: ACME.")
        table = doc.add_table(rows=1, cols=1)
        table.rows[0].cells[0].paragraphs[0].add_run("Proveedor ACME SRL")
        doc.sections[0].header.paragraphs[0].add_run("ACME - CONFIDENCIAL")
        doc.sections[0].footer.paragraphs[0].add_run("Página de ACME")
        raw = _doc_to_bytes(doc)

        result = await engine.apply_variable_mappings(
            raw, [{"text": "ACME", "variable": "company"}]
        )

        texts = _all_texts(result)
        joined = "\n".join(texts)
        assert "ACME" not in joined
        assert joined.count("{{ company }}") == 5
        # One per location
        assert any(
            "Cliente: {{ company }} en el cuerpo. Repetido: {{ company }}." == t
            for t in texts
        )
        assert "Proveedor {{ company }} SRL" in texts
        assert "{{ company }} - CONFIDENCIAL" in texts
        assert "Página de {{ company }}" in texts

    async def test_replacement_is_case_sensitive(self):
        engine = DocxTemplateEngine()

        doc = Document()
        doc.add_paragraph("acme y ACME no son lo mismo")
        raw = _doc_to_bytes(doc)

        result = await engine.apply_variable_mappings(
            raw, [{"text": "ACME", "variable": "company"}]
        )

        assert _body_texts(result)[0] == "acme y {{ company }} no son lo mismo"


class TestLongestFirst:
    async def test_longer_text_wins_over_contained_shorter_text(self):
        engine = DocxTemplateEngine()

        doc = Document()
        doc.add_paragraph("El Sr. Juan Pérez García y su hermano Juan Pérez firman.")
        raw = _doc_to_bytes(doc)

        result = await engine.apply_variable_mappings(
            raw,
            [
                # Deliberately shorter-first in the input: engine must reorder
                {"text": "Juan Pérez", "variable": "brother"},
                {"text": "Juan Pérez García", "variable": "main_party"},
            ],
        )

        assert (
            _body_texts(result)[0]
            == "El Sr. {{ main_party }} y su hermano {{ brother }} firman."
        )

    async def test_same_variable_for_two_texts_is_allowed(self):
        engine = DocxTemplateEngine()

        doc = Document()
        doc.add_paragraph("JUAN PÉREZ alias JP")
        raw = _doc_to_bytes(doc)

        result = await engine.apply_variable_mappings(
            raw,
            [
                {"text": "JUAN PÉREZ", "variable": "client_name"},
                {"text": "JP", "variable": "client_name"},
            ],
        )

        assert _body_texts(result)[0] == "{{ client_name }} alias {{ client_name }}"


# ---------------------------------------------------------------------------
# Validation and error cases
# ---------------------------------------------------------------------------


class TestValidation:
    @pytest.fixture
    def raw(self) -> bytes:
        doc = Document()
        doc.add_paragraph("Hola mundo")
        return _doc_to_bytes(doc)

    async def test_missing_text_raises_with_all_missing_texts(self, raw):
        engine = DocxTemplateEngine()

        with pytest.raises(MappingTextNotFoundError) as exc_info:
            await engine.apply_variable_mappings(
                raw,
                [
                    {"text": "NO EXISTE", "variable": "uno"},
                    {"text": "mundo", "variable": "dos"},
                    {"text": "TAMPOCO EXISTE", "variable": "tres"},
                ],
            )

        assert exc_info.value.missing_texts == ["NO EXISTE", "TAMPOCO EXISTE"]
        # The message must mention every missing text
        assert "NO EXISTE" in str(exc_info.value)
        assert "TAMPOCO EXISTE" in str(exc_info.value)

    async def test_empty_mappings_rejected(self, raw):
        engine = DocxTemplateEngine()
        with pytest.raises(InvalidVariableMappingError):
            await engine.apply_variable_mappings(raw, [])

    async def test_blank_text_rejected(self, raw):
        engine = DocxTemplateEngine()
        with pytest.raises(InvalidVariableMappingError):
            await engine.apply_variable_mappings(
                raw, [{"text": "   ", "variable": "valid_name"}]
            )

    @pytest.mark.parametrize(
        "bad_variable",
        ["Nombre", "clientName", "1abc", "name-x", "ñame", "with space", ""],
    )
    async def test_bad_variable_names_rejected(self, raw, bad_variable):
        engine = DocxTemplateEngine()
        with pytest.raises(InvalidVariableMappingError):
            await engine.apply_variable_mappings(
                raw, [{"text": "mundo", "variable": bad_variable}]
            )

    async def test_duplicate_texts_rejected(self, raw):
        engine = DocxTemplateEngine()
        with pytest.raises(InvalidVariableMappingError):
            await engine.apply_variable_mappings(
                raw,
                [
                    {"text": "mundo", "variable": "uno"},
                    {"text": "mundo", "variable": "dos"},
                ],
            )


# ---------------------------------------------------------------------------
# Non-matching paragraphs must not be touched
# ---------------------------------------------------------------------------


class TestUntouchedParagraphs:
    async def test_paragraphs_without_matches_keep_their_runs(self):
        engine = DocxTemplateEngine()

        doc = Document()
        p_untouched = doc.add_paragraph()
        r_a = p_untouched.add_run("Sin ")
        r_a.bold = True
        r_b = p_untouched.add_run("cambios ")
        r_b.italic = True
        r_c = p_untouched.add_run("aquí")
        r_c.font.name = "Courier New"
        r_c.font.size = Pt(9)
        doc.add_paragraph("Objetivo JUAN presente")
        raw = _doc_to_bytes(doc)

        result = await engine.apply_variable_mappings(
            raw, [{"text": "JUAN", "variable": "client_name"}]
        )

        out = _load(result)
        p0 = out.paragraphs[0]
        # Run count, texts, and key formatting must be exactly as authored
        assert [r.text for r in p0.runs] == ["Sin ", "cambios ", "aquí"]
        assert p0.runs[0].bold is True
        assert p0.runs[1].italic is True
        assert p0.runs[2].font.name == "Courier New"
        assert p0.runs[2].font.size == Pt(9)
        # And the matching paragraph was rewritten
        assert out.paragraphs[1].text == "Objetivo {{ client_name }} presente"

    async def test_input_bytes_are_not_mutated(self):
        engine = DocxTemplateEngine()

        doc = Document()
        doc.add_paragraph("Documento de JUAN")
        raw = _doc_to_bytes(doc)

        result = await engine.apply_variable_mappings(
            raw, [{"text": "JUAN", "variable": "client_name"}]
        )

        assert isinstance(result, bytes)
        assert result != raw
        # Re-parsing the ORIGINAL bytes still shows the original literal
        assert _body_texts(raw)[0] == "Documento de JUAN"


# ---------------------------------------------------------------------------
# ROUND-TRIP — the proof the whole feature works
# ---------------------------------------------------------------------------


class TestRoundTrip:
    async def test_example_to_template_to_rendered_document(self):
        """example docx → apply mappings → extract_variables finds exactly the
        mapped names → render with values → values appear in place."""
        engine = DocxTemplateEngine()

        doc = Document()
        # Literal split across two differently-formatted runs (Word reality)
        para = doc.add_paragraph()
        para.add_run("Entre ")
        r = para.add_run("JUAN ")
        r.bold = True
        para.add_run("PÉREZ")
        para.add_run(", en adelante EL CLIENTE.")
        # Literal inside a table cell
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].paragraphs[0].add_run("Domicilio:")
        table.rows[0].cells[1].paragraphs[0].add_run("CALLE FALSA 123")
        # Literal in the footer
        doc.sections[0].footer.paragraphs[0].add_run("Documento de ACME SRL")
        raw = _doc_to_bytes(doc)

        mappings = [
            {"text": "JUAN PÉREZ", "variable": "client_name"},
            {"text": "CALLE FALSA 123", "variable": "address"},
            {"text": "ACME SRL", "variable": "company"},
        ]

        template_bytes = await engine.apply_variable_mappings(raw, mappings)

        # 1. Extraction finds EXACTLY the mapped variable names
        variables_meta = await engine.extract_variables(template_bytes)
        names = {v["name"] for v in variables_meta}
        assert names == {"client_name", "address", "company"}

        # 2. Rendering puts the values in place
        rendered = await engine.render(
            template_bytes,
            {
                "client_name": "MARÍA LÓPEZ",
                "address": "AV. SIEMPRE VIVA 742",
                "company": "GLOBEX SA",
            },
        )
        texts = _all_texts(rendered)
        joined = "\n".join(texts)
        assert "Entre MARÍA LÓPEZ, en adelante EL CLIENTE." in texts
        assert "AV. SIEMPRE VIVA 742" in texts
        assert "Documento de GLOBEX SA" in texts
        # No placeholder and no original literal survived
        assert "{{" not in joined
        assert "JUAN" not in joined
        assert "CALLE FALSA" not in joined
        assert "ACME" not in joined
