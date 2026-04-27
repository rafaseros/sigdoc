"""
Unit tests for DocxTemplateEngine.extract_variables() — paragraph context behavior.

Key invariant: when a paragraph contains multiple {{ var }} placeholders, each
variable's ``contexts`` list must contain the SAME full paragraph string (stripped).
This allows the frontend to deduplicate by exact string match, collapsing N variables
in the same paragraph into a single rendered block with all placeholder pills.

Previously, _truncate_context() produced a different centered snippet per variable,
causing the frontend to see 2-3 distinct "snippets" for what is actually one paragraph.
"""
import io

import pytest
from docx import Document

from app.infrastructure.templating.docxtpl_engine import DocxTemplateEngine


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_docx_bytes(paragraphs: list[str]) -> bytes:
    """
    Build a minimal .docx in memory with one run per paragraph.

    ``paragraphs`` is a list of full paragraph text strings.
    """
    doc = Document()
    for text in paragraphs:
        para = doc.add_paragraph()
        para.add_run(text)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


class TestExtractVariablesContext:
    """extract_variables() must return the full paragraph as context, not a truncated snippet."""

    async def test_multi_var_paragraph_same_context_string(self):
        """
        A paragraph with multiple variables must produce the SAME context string
        for every variable in it. This is the core dedup-friendliness invariant.
        """
        engine = DocxTemplateEngine()
        # This paragraph is intentionally long (> 120 chars) so the old truncation
        # would have produced different snippets centered on each variable.
        para = (
            "La empresa {{ nombre_empresa }} con Matrícula de Comercio y NIT {{ numero }},"
            " con domicilio legal en la ciudad de {{ ciudad }}, representada legalmente."
        )
        raw = _make_docx_bytes([para])

        result = await engine.extract_variables(raw)

        # Must find all three variables
        names = {entry["name"] for entry in result}
        assert "nombre_empresa" in names
        assert "numero" in names
        assert "ciudad" in names

        # All three must have the SAME context string (full paragraph stripped)
        contexts_by_name = {entry["name"]: entry["contexts"] for entry in result}
        ctx_empresa = contexts_by_name["nombre_empresa"]
        ctx_numero = contexts_by_name["numero"]
        ctx_ciudad = contexts_by_name["ciudad"]

        assert len(ctx_empresa) == 1, f"Expected 1 context for nombre_empresa, got: {ctx_empresa}"
        assert len(ctx_numero) == 1, f"Expected 1 context for numero, got: {ctx_numero}"
        assert len(ctx_ciudad) == 1, f"Expected 1 context for ciudad, got: {ctx_ciudad}"

        assert ctx_empresa[0] == ctx_numero[0] == ctx_ciudad[0], (
            "All variables in the same paragraph must share the same context string. "
            f"nombre_empresa={ctx_empresa[0]!r}, "
            f"numero={ctx_numero[0]!r}, "
            f"ciudad={ctx_ciudad[0]!r}"
        )

    async def test_context_is_full_paragraph_not_truncated(self):
        """
        The context string must equal paragraph.strip() — no ellipsis markers,
        no mid-sentence cuts.
        """
        engine = DocxTemplateEngine()
        # Paragraph longer than the old 120-char window
        para = (
            "Este contrato es celebrado entre {{ parte_a }} y {{ parte_b }}, "
            "ambas partes con plena capacidad legal para contratar y obligarse "
            "conforme a las leyes vigentes del estado."
        )
        raw = _make_docx_bytes([para])

        result = await engine.extract_variables(raw)

        contexts_by_name = {entry["name"]: entry["contexts"] for entry in result}
        for var_name in ("parte_a", "parte_b"):
            ctx = contexts_by_name[var_name]
            assert len(ctx) == 1
            assert "..." not in ctx[0], (
                f"Context for '{var_name}' must not contain ellipsis markers. Got: {ctx[0]!r}"
            )
            assert ctx[0] == para.strip(), (
                f"Context for '{var_name}' must be the full paragraph stripped. "
                f"Expected: {para.strip()!r}, got: {ctx[0]!r}"
            )

    async def test_short_paragraph_returned_as_is(self):
        """Short paragraphs (under any length threshold) must also be returned verbatim."""
        engine = DocxTemplateEngine()
        para = "Hola {{ nombre }}, bienvenido."
        raw = _make_docx_bytes([para])

        result = await engine.extract_variables(raw)

        contexts_by_name = {entry["name"]: entry["contexts"] for entry in result}
        assert contexts_by_name["nombre"] == [para.strip()]

    async def test_context_strips_leading_trailing_whitespace(self):
        """Leading/trailing whitespace in the paragraph text is stripped."""
        engine = DocxTemplateEngine()
        para_with_ws = "   {{ nombre_cliente }} es el titular.   "
        raw = _make_docx_bytes([para_with_ws])

        result = await engine.extract_variables(raw)

        contexts_by_name = {entry["name"]: entry["contexts"] for entry in result}
        assert contexts_by_name["nombre_cliente"] == [para_with_ws.strip()]

    async def test_two_paragraphs_each_variable_gets_own_paragraph(self):
        """
        Variables in different paragraphs each get their own paragraph as context,
        not each other's.
        """
        engine = DocxTemplateEngine()
        para1 = "Firmado por {{ firmante }} el día de hoy."
        para2 = "El monto total es {{ monto }} pesos."
        raw = _make_docx_bytes([para1, para2])

        result = await engine.extract_variables(raw)

        contexts_by_name = {entry["name"]: entry["contexts"] for entry in result}
        assert contexts_by_name["firmante"] == [para1.strip()]
        assert contexts_by_name["monto"] == [para2.strip()]

    async def test_same_variable_in_two_paragraphs_gets_both_contexts(self):
        """
        A variable that appears in two different paragraphs must have two distinct
        context strings — one per paragraph.
        """
        engine = DocxTemplateEngine()
        para1 = "El cliente {{ nombre }} acepta los términos."
        para2 = "Notificar a {{ nombre }} en caso de cambios."
        raw = _make_docx_bytes([para1, para2])

        result = await engine.extract_variables(raw)

        contexts_by_name = {entry["name"]: entry["contexts"] for entry in result}
        ctx = contexts_by_name["nombre"]
        assert len(ctx) == 2, f"Expected 2 contexts for 'nombre', got: {ctx}"
        assert para1.strip() in ctx
        assert para2.strip() in ctx
